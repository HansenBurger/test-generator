"""调试文档结构"""
import sys
sys.path.insert(0, '.')
from docx import Document
import re

doc = Document('../tmp/新一代信贷系统建设项目_管理特色互联网贷款账单用例_贷款核算组_V0.9_20251028.docx')

print("=" * 60)
print("查找组件'个人贷款'之后的段落:")
print("=" * 60)

found_component = False
for i, p in enumerate(doc.paragraphs[25:40]):
    text = p.text.strip()
    if '个人贷款' in text and 'A阶段、B阶段' in text:
        found_component = True
        print(f'\n找到组件: 段落{i+25}: {text}')
    
    if found_component:
        # 检查是否有编号
        num_match = re.match(r'^(\d+(?:\.\d+)*)\.', text)
        if num_match:
            level = len(num_match.group(1).split('.'))
            print(f'段落{i+25}: {text[:80]} (层级: {level})')
        elif text and ('*' in text or len(text) > 0):
            print(f'段落{i+25}: {text[:80]} (无编号)')

print("\n" + "=" * 60)
print("查找所有包含编号和*的段落 (25-35):")
print("=" * 60)
for i, p in enumerate(doc.paragraphs[25:35]):
    text = p.text.strip()
    if text and (text[0].isdigit() or '*' in text):
        num_match = re.match(r'^(\d+(?:\.\d+)*)\.', text)
        level = len(num_match.group(1).split('.')) if num_match else 0
        print(f'段落{i+25}: {text[:100]} (层级: {level})')

