"""
Word文档解析服务 - 银行需求文档专用解析器
"""
import re
import os
import tempfile
from pathlib import Path
from typing import List, Optional, Dict, Tuple
from docx import Document
from docx.document import Document as DocumentType
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P

from app.models.schemas import (
    ParsedDocument, RequirementInfo, ActivityInfo, ComponentInfo,
    TaskInfo, StepInfo, InputElement, OutputElement, FunctionInfo
)
from app.utils.logger import parser_logger


class DocumentParser:
    """文档解析器 - 针对银行需求文档格式"""
    
    def __init__(self, doc_path: str):
        self._temp_docx_path = None  # 用于存储临时转换的 .docx 文件路径
        self.doc_path = doc_path
        parser_logger.info(f"开始初始化文档解析器，文件路径: {doc_path}")
        
        actual_doc_path = self._handle_doc_file(doc_path)
        
        try:
            self.doc = Document(actual_doc_path)
            self.paragraphs = [p for p in self.doc.paragraphs]
            self.tables = self.doc.tables
            self.used_tables = set()  # 记录已使用的表格索引，避免重复使用
            # 添加结果缓存，避免重复处理
            self._extraction_cache = {}  # 缓存提取结果: {(start_index, step_number): (input_elements, output_elements)}
            # 性能优化：文档长度 > 1000段落时启用优化模式
            self._optimization_mode = len(self.paragraphs) > 1000
            parser_logger.info(
                f"文档加载成功 - 段落数: {len(self.paragraphs)}, 表格数: {len(self.tables)}, "
                f"优化模式: {self._optimization_mode}"
            )
        except Exception as e:
            parser_logger.error(f"文档加载失败: {str(e)}", exc_info=True)
            # 清理临时文件
            self._cleanup_temp_file()
            raise
    
    def _handle_doc_file(self, doc_path: str) -> str:
        """处理 .doc 文件，如果是 .doc 格式则转换为 .docx"""
        doc_path_obj = Path(doc_path)
        
        # 如果已经是 .docx 格式，直接返回
        if doc_path_obj.suffix.lower() == '.docx':
            parser_logger.debug(f"文档格式为.docx，直接使用: {doc_path}")
            return doc_path
        
        # 如果是 .doc 格式，需要转换
        if doc_path_obj.suffix.lower() == '.doc':
            parser_logger.info(f"检测到.doc格式文件，开始转换为.docx: {doc_path}")
            converted_path = self._convert_doc_to_docx(doc_path)
            parser_logger.info(f"文档转换完成: {converted_path}")
            return converted_path
        
        # 其他格式，尝试直接打开（可能会失败）
        parser_logger.warning(f"未知文件格式，尝试直接打开: {doc_path}")
        return doc_path
    
    def _convert_doc_to_docx(self, doc_path: str) -> str:
        """将 .doc 文件转换为 .docx 格式
        
        在Windows上使用pywin32 + Microsoft Word COM接口
        在Linux上使用LibreOffice命令行工具
        """
        import platform
        import subprocess
        
        # 创建临时 .docx 文件
        temp_dir = tempfile.gettempdir()
        # 使用安全的文件名（移除特殊字符，避免路径问题）
        safe_filename = re.sub(r'[<>:"/\\|?*]', '_', os.path.basename(doc_path))
        # 移除原扩展名，添加.docx
        base_name = os.path.splitext(safe_filename)[0]
        temp_docx_path = os.path.join(
            temp_dir,
            f"converted_{base_name}.docx"
        )
        
        # 根据操作系统选择转换方式
        if platform.system() == "Windows":
            return self._convert_doc_to_docx_windows(doc_path, temp_docx_path)
        else:
            return self._convert_doc_to_docx_linux(doc_path, temp_docx_path)
    
    def _convert_doc_to_docx_windows(self, doc_path: str, output_path: str) -> str:
        """Windows下使用pywin32转换.doc文件"""
        try:
            import win32com.client
        except ImportError:
            raise ValueError(
                "无法处理 .doc 格式文件：需要安装 pywin32 库。"
                "请运行: pip install pywin32"
            )
        
        try:
            # 使用 Word COM 接口转换
            word_app = win32com.client.Dispatch("Word.Application")
            word_app.Visible = False
            word_app.DisplayAlerts = False
            
            try:
                # 打开 .doc 文件
                doc = word_app.Documents.Open(os.path.abspath(doc_path))
                
                # 保存为 .docx 格式
                doc.SaveAs2(
                    FileName=os.path.abspath(output_path),
                    FileFormat=16  # wdFormatXMLDocument = 16 (.docx)
                )
                
                doc.Close()
                word_app.Quit()
                
                # 保存临时文件路径，用于后续清理
                self._temp_docx_path = output_path
                
                return output_path
                
            except Exception as e:
                try:
                    word_app.Quit()
                except:
                    pass
                raise ValueError(
                    f"无法将 .doc 文件转换为 .docx 格式：{str(e)}。"
                    "请确保已安装 Microsoft Word，或手动将文件转换为 .docx 格式。"
                )
        except Exception as e:
            if isinstance(e, ValueError):
                raise
            raise ValueError(
                f"无法处理 .doc 格式文件：{str(e)}。"
                "请确保已安装 Microsoft Word，或手动将文件转换为 .docx 格式。"
            )
    
    def _convert_doc_to_docx_linux(self, doc_path: str, output_path: str) -> str:
        """Linux下转换.doc文件
        
        优先使用antiword + pandoc（更快），失败则回退到LibreOffice
        """
        # 首先尝试使用antiword + pandoc（更快）
        try:
            return self._convert_doc_to_docx_antiword_pandoc(doc_path, output_path)
        except Exception as e:
            parser_logger.warning(f"antiword+pandoc转换失败，回退到LibreOffice: {str(e)}")
            # 回退到LibreOffice
            return self._convert_doc_to_docx_libreoffice(doc_path, output_path)
    
    def _convert_doc_to_docx_antiword_pandoc(self, doc_path: str, output_path: str) -> str:
        """使用antiword + pandoc转换（更快，保留章节和表格结构）"""
        import subprocess
        
        try:
            # 确保输出目录存在
            output_dir = os.path.dirname(output_path)
            os.makedirs(output_dir, exist_ok=True)
            
            # 方法：使用antiword提取文本（保留表格结构），然后pandoc转换为docx
            # antiword可以保留基本的表格结构（通过制表符）
            parser_logger.info(f"使用antiword提取文本: {doc_path}")
            
            # 使用antiword提取文本（-m参数保留表格格式）
            txt_path = output_path.replace('.docx', '.txt')
            antiword_result = subprocess.run(
                ["antiword", "-m", "UTF-8.txt", doc_path],  # -m指定映射文件，UTF-8.txt保留格式
                capture_output=True,
                text=True,
                timeout=60
            )
            
            if antiword_result.returncode != 0:
                # 如果-m参数失败，尝试不带参数
                parser_logger.warning("antiword -m参数失败，尝试普通模式")
                antiword_result = subprocess.run(
                    ["antiword", doc_path],
                    capture_output=True,
                    text=True,
                    timeout=60
                )
                
                if antiword_result.returncode != 0:
                    raise ValueError(f"antiword提取失败: {antiword_result.stderr}")
            
            # 保存文本（保留表格的制表符结构）
            with open(txt_path, 'w', encoding='utf-8') as f:
                f.write(antiword_result.stdout)
            
            # 步骤2: 使用pandoc将文本转换为docx（保留表格）
            parser_logger.info(f"使用pandoc转换文本到docx: {txt_path}")
            pandoc_result = subprocess.run(
                ["pandoc", "-f", "plain", "-t", "docx", "--wrap=none", "-o", output_path, txt_path],
                capture_output=True,
                text=True,
                timeout=60
            )
            
            # 清理临时文件
            if os.path.exists(txt_path):
                try:
                    os.unlink(txt_path)
                except:
                    pass
            
            if pandoc_result.returncode != 0:
                raise ValueError(f"pandoc转换失败: {pandoc_result.stderr}")
            
            # 验证输出文件
            if not os.path.exists(output_path) or os.path.getsize(output_path) == 0:
                raise ValueError("转换后的.docx文件为空或不存在")
            
            parser_logger.info(f"antiword+pandoc转换成功: {output_path}")
            self._temp_docx_path = output_path
            return output_path
            
        except FileNotFoundError as e:
            raise ValueError(
                "无法处理 .doc 格式文件：未找到antiword或pandoc。"
                "请确保已安装antiword和pandoc（在Docker容器中应已自动安装）。"
            )
        except subprocess.TimeoutExpired:
            raise ValueError(
                "转换.doc文件超时，文件可能过大或已损坏。"
                "请尝试手动将文件转换为.docx格式。"
            )
        except Exception as e:
            raise ValueError(f"antiword+pandoc转换失败: {str(e)}")
    
    def _convert_doc_to_docx_libreoffice(self, doc_path: str, output_path: str) -> str:
        """使用LibreOffice转换（备选方案）"""
        return self._convert_doc_to_docx_cli(doc_path, output_path)
    
    def _convert_doc_to_docx_cli(self, doc_path: str, output_path: str) -> str:
        """使用命令行模式转换（备选方案）"""
        import subprocess
        
        try:
            # 获取输出目录
            output_dir = os.path.dirname(output_path)
            
            # 确保输出目录存在
            os.makedirs(output_dir, exist_ok=True)
            
            # 使用soffice批处理模式
            # 使用优化的批处理模式，移除--safe-mode（可能反而更慢）
            # 使用最少的参数，减少初始化时间
            cmd = [
                "soffice",
                "--headless",
                "--invisible",
                "--nodefault",
                "--nolockcheck",
                "--nologo",
                "--norestore",
                "--convert-to", "docx:MS Word 2007 XML",  # 明确指定输出格式
                "--outdir", output_dir,
                doc_path
            ]
            
            # 优化环境变量，减少LibreOffice初始化时间
            env = dict(os.environ)
            env.update({
                "HOME": "/tmp",
                "SAL_USE_VCLPLUGIN": "headless",  # 强制使用headless插件
                "SAL_DISABLE_OPENCL": "1",  # 禁用OpenCL加速（减少初始化）
                "SAL_DISABLE_OPENCL_CLEANUP": "1",  # 禁用OpenCL清理
                "OOO_DISABLE_RECOVERY": "1",  # 禁用恢复功能
            })
            
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=300,  # 300秒超时
                env=env,
                preexec_fn=lambda: os.nice(10) if hasattr(os, 'nice') else None  # 降低优先级，避免阻塞
            )
            
            # LibreOffice即使成功也可能返回非0退出码，所以主要检查文件是否生成
            # 但如果有明显的错误信息，还是抛出异常
            if result.returncode != 0 and "error" in result.stderr.lower():
                raise subprocess.CalledProcessError(
                    result.returncode,
                    cmd,
                    result.stderr
                )
            
            # LibreOffice输出的文件名是基于输入文件名生成的（移除扩展名后加.docx）
            input_basename = os.path.splitext(os.path.basename(doc_path))[0]
            # 清理文件名中的特殊字符（LibreOffice可能会处理）
            safe_basename = re.sub(r'[<>:"/\\|?*]', '_', input_basename)
            
            # 尝试查找实际生成的文件（可能有多种变体）
            possible_files = [
                os.path.join(output_dir, f"{input_basename}.docx"),
                os.path.join(output_dir, f"{safe_basename}.docx"),
                os.path.join(output_dir, os.path.basename(doc_path).replace('.doc', '.docx').replace('.DOC', '.docx')),
            ]
            
            # 如果都不存在，列出目录中所有.docx文件
            generated_file = None
            for possible_file in possible_files:
                if os.path.exists(possible_file):
                    generated_file = possible_file
                    break
            
            # 如果还是没找到，搜索输出目录中的所有.docx文件
            if not generated_file:
                for file in os.listdir(output_dir):
                    if file.endswith('.docx') and file.startswith(input_basename[:10]):  # 至少前10个字符匹配
                        generated_file = os.path.join(output_dir, file)
                        break
            
            if not generated_file:
                # 最后尝试：查找最近生成的.docx文件
                docx_files = [f for f in os.listdir(output_dir) if f.endswith('.docx')]
                if docx_files:
                    # 按修改时间排序，取最新的
                    docx_files.sort(key=lambda f: os.path.getmtime(os.path.join(output_dir, f)), reverse=True)
                    generated_file = os.path.join(output_dir, docx_files[0])
            
            if not generated_file or not os.path.exists(generated_file):
                error_msg = f"LibreOffice转换完成，但未找到生成的.docx文件。"
                if result.stderr:
                    error_msg += f" 错误信息: {result.stderr[:200]}"
                raise ValueError(error_msg)
            
            # 如果生成的文件名与期望的不同，重命名
            if generated_file != output_path:
                if os.path.exists(output_path):
                    os.unlink(output_path)
                os.rename(generated_file, output_path)
            
            # 验证文件确实存在且不为空
            if not os.path.exists(output_path) or os.path.getsize(output_path) == 0:
                raise ValueError("转换后的.docx文件为空或不存在")
            
            # 保存临时文件路径，用于后续清理
            self._temp_docx_path = output_path
            
            return output_path
            
        except subprocess.TimeoutExpired:
            raise ValueError(
                "转换.doc文件超时，文件可能过大或已损坏。"
                "请尝试手动将文件转换为.docx格式。"
            )
        except FileNotFoundError:
            raise ValueError(
                "无法处理 .doc 格式文件：未找到LibreOffice。"
                "请确保已安装LibreOffice（在Docker容器中应已自动安装）。"
            )
        except subprocess.CalledProcessError as e:
            raise ValueError(
                f"无法将 .doc 文件转换为 .docx 格式：{e.stderr or str(e)}。"
                "请确保LibreOffice已正确安装，或手动将文件转换为 .docx 格式。"
            )
        except Exception as e:
            raise ValueError(
                f"无法处理 .doc 格式文件：{str(e)}。"
                "请确保LibreOffice已正确安装，或手动将文件转换为 .docx 格式。"
            )
    
    def _cleanup_temp_file(self):
        """清理临时转换的 .docx 文件"""
        if self._temp_docx_path and os.path.exists(self._temp_docx_path):
            try:
                os.unlink(self._temp_docx_path)
            except:
                pass
            self._temp_docx_path = None
    
    def __del__(self):
        """析构函数，清理临时文件"""
        self._cleanup_temp_file()
    
    def parse(self) -> ParsedDocument:
        """解析文档主方法"""
        parser_logger.info(f"开始解析文档: {self.doc_path}")
        
        # 1. 识别文档类型
        doc_type = self._identify_document_type()
        parser_logger.info(f"识别文档类型: {doc_type}")
        
        try:
            if doc_type == "modeling":
                parser_logger.info("开始解析建模需求文档")
                result = self._parse_modeling_document()
                parser_logger.info(
                    f"建模需求文档解析成功 - 用例名称: {result.requirement_info.case_name if result.requirement_info else 'N/A'}"
                )
                return result
            elif doc_type == "non_modeling":
                parser_logger.info("开始解析非建模需求文档")
                result = self._parse_non_modeling_document()
                parser_logger.info(
                    f"非建模需求文档解析成功 - 需求名称: {result.requirement_name or 'N/A'}, "
                    f"功能数: {len(result.functions) if result.functions else 0}"
                )
                return result
            else:
                error_msg = "无法识别文档类型：未找到'用例版本控制信息'表或'文件受控信息'/'文档受控信息'表"
                parser_logger.error(error_msg)
                raise ValueError(error_msg)
        except Exception as e:
            parser_logger.error(f"文档解析失败: {str(e)}", exc_info=True)
            raise
    
    def _identify_document_type(self) -> Optional[str]:
        """识别文档类型：建模需求或非建模需求"""
        # 优先级1：查找"用例版本控制信息"（建模需求的明确标识）
        for para in self.paragraphs[:100]:
            text = para.text.strip()
            if "用例版本控制信息" in text:
                # 检查是否有包含"版本"字段的表格
                for table in self.tables:
                    if len(table.rows) < 1:
                        continue
                    header_row = table.rows[0]
                    header_text = ' '.join([cell.text.strip() for cell in header_row.cells])
                    if "版本" in header_text:
                        return "modeling"
        
        # 优先级2：查找"文件受控信息"或"文档受控信息"（非建模需求的明确标识）
        # 同时检查是否有"功能清单"（非建模需求的另一个特征）
        has_file_control = False
        has_function_list = False
        
        for para in self.paragraphs[:100]:
            text = para.text.strip()
            if "文件受控信息" in text or "文档受控信息" in text:
                has_file_control = True
            if "功能清单" in text:
                has_function_list = True
        
        # 检查表格
        for table in self.tables:
            if len(table.rows) < 1:
                continue
            header_row = table.rows[0]
            header_text = ' '.join([cell.text.strip() for cell in header_row.cells])
            
            if ("文件编号" in header_text or "文件名称" in header_text or 
                "文档受控信息" in header_text):
                has_file_control = True
            
            if "业务功能名称" in header_text or "功能名称" in header_text:
                has_function_list = True
        
        # 如果有文件受控信息或功能清单，识别为非建模需求
        if has_file_control or has_function_list:
            return "non_modeling"
        
        # 优先级3：查找"版本控制信息"（不带"用例"前缀，可能是建模需求）
        # 但需要更严格的判断：必须同时有"需求用例概述"
        has_version_control = False
        has_requirement_overview = False
        
        for para in self.paragraphs[:100]:
            text = para.text.strip()
            if "版本控制信息" in text and "用例" not in text:
                has_version_control = True
            if "需求用例概述" in text:
                has_requirement_overview = True
        
        # 检查是否有包含"版本"字段的表格
        for table in self.tables:
            if len(table.rows) < 1:
                continue
            header_row = table.rows[0]
            header_text = ' '.join([cell.text.strip() for cell in header_row.cells])
            if "版本" in header_text and has_version_control:
                # 如果同时有"需求用例概述"，才识别为建模需求
                if has_requirement_overview:
                    return "modeling"
        
        return None
    
    def _parse_modeling_document(self) -> ParsedDocument:
        """解析建模需求文档"""
        # 1. 提取版本编号
        version = self._extract_version()
        if not version:
            raise ValueError("无法提取版本信息：未找到'用例版本控制信息'表或表中无版本数据")
        
        # 2. 提取需求用例基本信息
        requirement_info = self._extract_requirement_info()
        if not requirement_info.case_name:
            raise ValueError("无法提取需求基本信息：未找到'需求用例概述'表或表中无用例名称")
        
        # 3. 提取活动名称（从任务设计部分）
        activity_name = self._extract_activity_name()
        
        # 4. 提取组件、任务、步骤信息（从任务规则说明部分）
        components = self._extract_all_components()
        
        # 构建活动信息（组件信息放在活动下，但XMind生成时会分别处理）
        activity = ActivityInfo(
            name=activity_name or "",
            components=components
        )
        
        return ParsedDocument(
            version=version,
            requirement_info=requirement_info,
            activities=[activity] if activity_name or components else [],
            document_type="modeling"
        )
    
    def _parse_non_modeling_document(self) -> ParsedDocument:
        """解析非建模需求文档"""
        # 1. 提取文件编号和文件名称
        file_number, file_name = self._extract_file_controlled_info()
        
        # 2. 提取需求名称
        requirement_name = self._extract_requirement_name(file_name)
        
        # 3. 提取设计者
        designer = self._extract_designer()
        
        # 4. 提取功能列表
        functions = self._extract_functions()
        
        if not functions:
            raise ValueError("无法提取功能列表：未找到'功能清单'表或表中无功能数据")
        
        # 构建需求基本信息（使用需求名称作为用例名称）
        requirement_info = RequirementInfo(case_name=requirement_name or "")
        
        return ParsedDocument(
            version=file_number or "",  # 使用文件编号作为版本
            requirement_info=requirement_info,
            activities=[],
            document_type="non_modeling",
            file_number=file_number,
            file_name=file_name,
            requirement_name=requirement_name,
            designer=designer,
            functions=functions
        )
    
    def _validate_document(self) -> bool:
        """验证文档是否包含用例版本控制信息表"""
        # 检查前50个段落中是否包含"用例版本控制信息"
        for para in self.paragraphs[:50]:
            text = para.text.strip()
            if "用例版本控制信息" in text:
                # 检查是否有包含"版本"字段的表格
                for table in self.tables:
                    if len(table.rows) < 1:
                        continue
                    header_row = table.rows[0]
                    header_text = ' '.join([cell.text.strip() for cell in header_row.cells])
                    if "版本" in header_text:
                        return True
        return False
    
    def _extract_version(self) -> str:
        """从用例版本控制信息表提取版本编号"""
        for table in self.tables:
            if len(table.rows) < 2:  # 至少要有表头和数据行
                continue
            
            # 检查表头是否包含"版本"字段
            header_row = table.rows[0]
            header_text = ' '.join([cell.text.strip() for cell in header_row.cells])
            
            if "版本" in header_text:
                # 找到版本列的索引（通常是第一列）
                version_col_idx = 0
                for idx, cell in enumerate(header_row.cells):
                    if "版本" in cell.text.strip():
                        version_col_idx = idx
                        break
                
                # 取最后非空行的版本列值
                for row in reversed(table.rows[1:]):  # 跳过表头
                    if len(row.cells) > version_col_idx:
                        version = row.cells[version_col_idx].text.strip()
                        if version:
                            return version
        return ""
    
    def _extract_requirement_info(self) -> RequirementInfo:
        """提取需求用例基本信息"""
        info = RequirementInfo(case_name="")
        # 查找包含"需求用例概述"的段落，然后查找其后紧邻的表格
        found_overview = False
        for block in self._iter_block_items():
            if isinstance(block, Paragraph):
                text = block.text.strip()
                if "需求用例概述" in text:
                    found_overview = True
                continue
            
            if found_overview and isinstance(block, Table):
                header_row_idx = self._find_requirement_header_row(block)
                if header_row_idx is not None:
                    self._parse_requirement_table(block, info, header_row_idx)
                    return info
        
        # 回退：如果未按顺序找到，扫描所有表格
        for table in self.tables:
            header_row_idx = self._find_requirement_header_row(table)
            if header_row_idx is not None:
                self._parse_requirement_table(table, info, header_row_idx)
                break
        
        return info

    def _iter_block_items(self):
        """按文档顺序迭代段落和表格"""
        for child in self.doc.element.body.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, self.doc)
            elif isinstance(child, CT_Tbl):
                yield Table(child, self.doc)

    def _find_requirement_header_row(self, table: Table) -> Optional[int]:
        """查找包含'用例名称'的表头行索引"""
        for idx, row in enumerate(table.rows):
            row_text = " ".join(cell.text.strip() for cell in row.cells)
            if "用例名称" in row_text:
                return idx
        return None
    
    def _parse_requirement_table(self, table: Table, info: RequirementInfo, header_row_idx: int = 0):
        """解析需求用例概述表格"""
        if len(table.rows) < 1 or header_row_idx >= len(table.rows):
            return
        
        # 获取表头
        header_row = table.rows[header_row_idx]
        headers = [cell.text.strip() for cell in header_row.cells]
        
        # 特殊处理：如果表头中"用例名称"后面直接跟着值（如：['用例名称', '管理特色互联网贷款账单', ...]）
        case_name_idx = self._find_column_index(headers, ["用例名称"])
        if case_name_idx >= 0 and case_name_idx + 1 < len(headers):
            # 检查用例名称列后面的列是否是值（不是其他字段名）
            next_cell = headers[case_name_idx + 1]
            if next_cell and next_cell != '/' and "用例" not in next_cell and "名称" not in next_cell:
                # 这可能是用例名称的值
                if not info.case_name:
                    info.case_name = next_cell
        
        # 查找各字段的列索引
        channel_idx = self._find_column_index(headers, ["渠道（C）", "渠道"])
        product_idx = self._find_column_index(headers, ["产品（P）", "P产品（P）", "产品"])
        customer_idx = self._find_column_index(headers, ["客户（C）", "客户"])
        partner_idx = self._find_column_index(headers, ["合作方（P）", "P合作方（P）", "合作方"])
        
        # 解析数据行（可能是横向布局：第一行是表头，第二行是值）
        value_row_idx = header_row_idx + 1
        if len(table.rows) > value_row_idx:
            value_row = table.rows[value_row_idx]
            values = [cell.text.strip() for cell in value_row.cells]
            
            if channel_idx >= 0 and channel_idx < len(values):
                value = values[channel_idx]
                if value and value != '/':
                    info.channel = value
            
            if product_idx >= 0 and product_idx < len(values):
                value = values[product_idx]
                if value and value != '/':
                    info.product = value
            
            if customer_idx >= 0 and customer_idx < len(values):
                value = values[customer_idx]
                if value and value != '/':
                    info.customer = value
            
            if partner_idx >= 0 and partner_idx < len(values):
                value = values[partner_idx]
                if value and value != '/':
                    info.partner = value
        
        # 也尝试纵向布局：第一列是键，第二列是值
        for row in table.rows[header_row_idx:]:
            if len(row.cells) >= 2:
                key = row.cells[0].text.strip()
                value = row.cells[1].text.strip()
                
                if value and value != '/':
                    if "用例名称" in key and not info.case_name:
                        info.case_name = value
                    elif ("渠道" in key and "（C）" in key) and not info.channel:
                        info.channel = value
                    elif (("产品" in key and "（P）" in key) or "P产品（P）" in key) and not info.product:
                        info.product = value
                    elif ("客户" in key and "（C）" in key) and not info.customer:
                        info.customer = value
                    elif (("合作方" in key and "（P）" in key) or "P合作方（P）" in key) and not info.partner:
                        info.partner = value
    
    def _extract_activity_name(self) -> Optional[str]:
        """提取活动名称：从'# 任务设计'部分提取第一个子标题"""
        for i, para in enumerate(self.paragraphs):
            text = para.text.strip()
            
            # 查找"任务设计"标题（一级标题）
            if "任务设计" in text:
                # 检查是否是标题样式（Heading 1）
                if self._is_heading(para, level=1):
                    # 查找下一个二级标题（##级别）
                    for j in range(i + 1, min(i + 50, len(self.paragraphs))):
                        next_para = self.paragraphs[j]
                        next_text = next_para.text.strip()
                        
                        # 如果遇到下一个一级标题，停止搜索
                        if self._is_heading(next_para, level=1) and "任务设计" not in next_text:
                            break
                        
                        # 检查是否是二级标题
                        if self._is_heading(next_para, level=2):
                            activity_name = self._strip_phase_suffix(next_text)
                            # 排除特定关键词
                            exclude_keywords = ["需求用例概述", "活动任务图", "业务流程图", 
                                               "任务设计", "业务步骤/功能描述", "规则说明",
                                               "任务清单", "任务流程图", "流程描述"]
                            if activity_name and activity_name not in exclude_keywords:
                                return activity_name
        
        return None
    
    def _extract_all_components(self) -> List[ComponentInfo]:
        """提取所有组件、任务、步骤信息：从'# 任务规则说明'部分提取"""
        components = []
        exclude_keywords = ["任务规则说明", "输入输出", "业务流程", "业务规则",
                           "页面控制", "数据验证", "前置条件", "后置条件",
                           "任务-业务步骤/功能清单", "业务步骤/功能描述", "规则说明",
                           "错误处理", "权限说明", "用户操作注释"]
        
        for i, para in enumerate(self.paragraphs):
            text = para.text.strip()
            
            # 查找"任务规则说明"标题（一级标题）
            if "任务规则说明" in text:
                if self._is_heading(para, level=1):
                    # 首先找到搜索的结束位置（下一个一级标题）
                    end_index = len(self.paragraphs)  # 默认到文档末尾
                    
                    for j in range(i + 1, len(self.paragraphs)):
                        next_para = self.paragraphs[j]
                        next_text = next_para.text.strip()
                        
                        # 如果遇到下一个一级标题，停止搜索
                        if self._is_heading(next_para, level=1) and "任务规则说明" not in next_text:
                            end_index = j
                            break
                    
                    # 在确定的范围内查找所有组件名称（二级标题：##级别）
                    for j in range(i + 1, end_index):
                        next_para = self.paragraphs[j]
                        next_text = next_para.text.strip()
                        
                        # 检查是否是二级标题（组件名称）
                        if self._is_heading(next_para, level=2):
                            component_name = self._strip_phase_suffix(next_text)
                            if component_name and component_name not in exclude_keywords:
                                # 提取该组件下的所有任务
                                tasks = self._extract_tasks(j + 1, component_name, exclude_keywords)
                                component = ComponentInfo(name=component_name, tasks=tasks)
                                components.append(component)
        
        return components
    
    def _extract_tasks(self, start_index: int, component_name: str, exclude_keywords: List[str]) -> List[TaskInfo]:
        """提取任务列表（从组件名称后开始）
        
        使用动态边界检测，自动找到下一个组件或一级标题作为结束位置
        """
        tasks = []
        current_task = None
        
        # 首先找到搜索的结束位置（下一个组件或一级标题）
        end_index = len(self.paragraphs)  # 默认到文档末尾
        
        for i in range(start_index + 1, len(self.paragraphs)):
            para = self.paragraphs[i]
            text = para.text.strip()
            
            # 如果遇到下一个二级标题（新的组件），停止搜索
            if self._is_heading(para, level=2) and component_name not in text:
                end_index = i
                break
            
            # 如果遇到一级标题，停止搜索
            if self._is_heading(para, level=1):
                end_index = i
                break
        
        # 在确定的范围内搜索任务
        for i in range(start_index, end_index):
            para = self.paragraphs[i]
            text = para.text.strip()
            
            # 检查是否是三级标题（任务名称）
            if self._is_heading(para, level=3):
                task_name = self._strip_phase_suffix(text)
                if task_name and task_name not in exclude_keywords:
                    # 保存上一个任务
                    if current_task:
                        tasks.append(current_task)
                    
                    # 创建新任务
                    current_task = TaskInfo(name=task_name, steps=[])
                    # 提取该任务下的所有步骤
                    steps = self._extract_steps(i + 1, task_name, exclude_keywords)
                    current_task.steps = steps
        
        # 添加最后一个任务
        if current_task:
            tasks.append(current_task)
        
        return tasks
    
    def _extract_steps(self, start_index: int, task_name: str, exclude_keywords: List[str]) -> List[StepInfo]:
        """提取步骤列表（从任务名称后开始）
        
        使用动态边界检测，自动找到下一个任务/组件/一级标题作为结束位置
        这样可以适应任意长度的文档，不需要固定搜索范围
        """
        steps = []
        
        # 首先找到搜索的结束位置（下一个任务、组件或一级标题）
        end_index = len(self.paragraphs)  # 默认到文档末尾
        
        for i in range(start_index + 1, len(self.paragraphs)):
            para = self.paragraphs[i]
            text = para.text.strip()
            
            # 如果遇到下一个三级标题（新的任务），停止搜索
            if self._is_heading(para, level=3) and task_name not in text:
                end_index = i
                break
            
            # 如果遇到二级标题（新的组件），停止搜索
            if self._is_heading(para, level=2):
                end_index = i
                break
            
            # 如果遇到一级标题，停止搜索
            if self._is_heading(para, level=1):
                end_index = i
                break
        
        # 在确定的范围内搜索步骤
        for i in range(start_index, end_index):
            para = self.paragraphs[i]
            text = para.text.strip()
            
            # 检查是否是四级标题（步骤名称）
            if self._is_heading(para, level=4):
                # 提取章节序号（如4.1.1.20）
                step_number_match = re.match(r"(\d+\.\d+\.\d+\.\d+)\.?\s*(.+)", text)
                if step_number_match:
                    step_number = step_number_match.group(1)  # 如"4.1.1.20"
                    step_name = self._strip_phase_suffix(step_number_match.group(2).strip())
                else:
                    # 如果没有章节序号，尝试只匹配步骤名称
                    step_number = None
                    step_name = self._strip_phase_suffix(text)
                    if not step_name:
                        continue
                
                if step_name not in exclude_keywords:
                    # 提取输入输出要素，传入步骤序号用于精确定位
                    input_elements, output_elements = self._extract_input_output_elements(i + 1, step_number)
                    step = StepInfo(
                        name=step_name,
                        input_elements=input_elements,
                        output_elements=output_elements
                    )
                    steps.append(step)
        
        return steps
    
    def _extract_input_output_elements(self, start_index: int, step_number: Optional[str] = None) -> Tuple[List[InputElement], List[OutputElement]]:
        """提取输入输出要素（优化版 - 支持文档末尾特殊处理和增强标记识别）
        
        Args:
            start_index: 步骤标题后的起始索引
            step_number: 步骤的章节序号（如"4.1.1.20"），用于精确定位输入输出子章节
        """
        # 检查缓存
        cache_key = (start_index, step_number)
        if cache_key in self._extraction_cache:
            parser_logger.debug(f"使用缓存结果 - 步骤序号: {step_number}, 起始位置: {start_index}")
            return self._extraction_cache[cache_key]
        
        input_elements = []
        output_elements = []
        
        # 判断是否为文档末尾的步骤（用于特殊处理）- 使用绝对位置计算
        is_document_end = start_index > len(self.paragraphs) * 0.7
        doc_length = len(self.paragraphs)
        
        parser_logger.info(
            f"提取输入输出要素 - 步骤序号: {step_number}, 起始位置: {start_index}/{doc_length} "
            f"({start_index/doc_length*100:.1f}%), 是否文档末尾: {is_document_end}"
        )
        
        # 策略1：如果有步骤序号，先查找对应的"输入输出"子章节（如4.1.1.20.2）
        input_output_index = -1
        input_output_found = False
        
        if step_number:
            # 文档末尾步骤：扩大搜索范围
            search_limit = min(start_index + 200, doc_length) if is_document_end else min(start_index + 100, doc_length)
            
            # 查找"输入输出"子章节，格式为：step_number.2（如4.1.1.20.2）
            input_output_patterns = [
                re.compile(rf"^{re.escape(step_number)}\.2\s*输入输出"),
                re.compile(rf"^{re.escape(step_number)}\.2\.\s*输入输出"),
                re.compile(rf"^{re.escape(step_number)}\.2\s+输入输出"),
            ]
            
            for i in range(start_index, search_limit):
                para = self.paragraphs[i]
                text = para.text.strip()
                
                # 检查是否是"输入输出"子章节
                for pattern in input_output_patterns:
                    if pattern.match(text):
                        if self._is_heading(para, level=5):
                            input_output_found = True
                            input_output_index = i
                            parser_logger.debug(f"找到输入输出子章节（精确匹配）: 索引 {i}, 文本: {repr(text)}")
                            break
                
                # 也检查简单的文本匹配
                if step_number + ".2" in text and "输入输出" in text:
                    if self._is_heading(para, level=5):
                        input_output_found = True
                        input_output_index = i
                        parser_logger.debug(f"找到输入输出子章节（文本匹配）: 索引 {i}, 文本: {repr(text)}")
                        break
                
                if input_output_found:
                    break
        
        # 策略2：如果没有找到或没有步骤序号，使用原来的方法
        if not input_output_found:
            search_limit = min(start_index + 100, doc_length) if is_document_end else min(start_index + 50, doc_length)
            for i in range(start_index, search_limit):
                para = self.paragraphs[i]
                text = para.text.strip()
                
                if "输入输出" in text:
                    if self._is_heading(para, level=5):
                        input_output_found = True
                        input_output_index = i
                        parser_logger.debug(f"找到输入输出子章节（通用匹配）: 索引 {i}, 文本: {repr(text)}")
                        break

        # 确定搜索范围 - 文档末尾使用绝对位置计算
        if input_output_found:
            # 如果找到了"输入输出"子章节，在该章节范围内搜索
            search_start = input_output_index + 1
            # 文档末尾：使用更大的搜索范围
            if is_document_end:
                search_end = min(input_output_index + 300, doc_length)  # 文档末尾允许更大范围
            else:
                search_end = min(input_output_index + 100, doc_length)
            
            # 查找结束位置（下一个四级标题、三级标题、二级标题或一级标题）
            for i in range(search_start, search_end):
                para = self.paragraphs[i]
                if (self._is_heading(para, level=4) or 
                    self._is_heading(para, level=3) or 
                    self._is_heading(para, level=2) or 
                    self._is_heading(para, level=1)):
                    search_end = i
                    break
            
            # 文档末尾特殊处理：如果没有找到下一个标题，扩展到文档末尾
            if is_document_end and search_end < doc_length * 0.95:
                search_end = doc_length
        else:
            # 如果没有找到"输入输出"子章节，回退到原来的逻辑：在步骤范围内搜索
            search_start = start_index
            # 文档末尾：使用更大的搜索范围
            if is_document_end:
                search_end = min(start_index + 300, doc_length)  # 文档末尾允许更大范围
            else:
                search_end = min(start_index + 100, doc_length)
            
            # 查找结束位置（下一个四级标题、三级标题、二级标题或一级标题）
            for i in range(search_start, search_end):
                para = self.paragraphs[i]
                if (self._is_heading(para, level=4) or 
                    self._is_heading(para, level=3) or 
                    self._is_heading(para, level=2) or 
                    self._is_heading(para, level=1)):
                    search_end = i
                    break
            
            # 文档末尾特殊处理：如果没有找到下一个标题，扩展到文档末尾
            if is_document_end and search_end < doc_length * 0.95:
                search_end = doc_length
        
        parser_logger.debug(
            f"搜索范围确定 - 起始: {search_start}, 结束: {search_end}, "
            f"范围大小: {search_end - search_start}"
        )
        
        # 在搜索范围内查找"输入要素"和"输出要素"文本 - 支持多种变体格式
        input_marker_variants = ["输入要素", "输入要素：", "输入要素表", "输入要素说明", 
                                "输入", "输入字段", "输入参数"]
        output_marker_variants = ["输出要素", "输出要素：", "输出要素表", "输出要素说明",
                                 "输出", "输出字段", "输出参数"]
        
        found_input_text = False
        found_output_text = False
        input_not_involved = False
        output_not_involved = False
        input_marker_index = -1
        output_marker_index = -1
        
        for i in range(search_start, search_end):
            para = self.paragraphs[i]
            text = para.text.strip()
            
            # 查找"输入要素"文本 - 支持多种变体
            if not found_input_text:
                for variant in input_marker_variants:
                    if variant in text:
                        # 排除包含"输出"的情况（避免误匹配）
                        if "输出" not in text or ("输入" in text and text.index("输入") < text.index("输出")):
                            found_input_text = True
                            input_marker_index = i
                            parser_logger.debug(f"找到输入要素标记: 索引 {i}, 文本: {repr(text)}")
                            break
                
                # 也支持"输入"和"要素"分开的情况
                if not found_input_text and "输入" in text and "要素" in text:
                    found_input_text = True
                    input_marker_index = i
                    parser_logger.debug(f"找到输入要素标记（分离匹配）: 索引 {i}, 文本: {repr(text)}")
            
            # 查找"输出要素"文本 - 支持多种变体
            if not found_output_text:
                for variant in output_marker_variants:
                    if variant in text:
                        found_output_text = True
                        output_marker_index = i
                        parser_logger.debug(f"找到输出要素标记: 索引 {i}, 文本: {repr(text)}")
                        break
                
                # 也支持"输出"和"要素"分开的情况
                if not found_output_text and "输出" in text and "要素" in text:
                    found_output_text = True
                    output_marker_index = i
                    parser_logger.debug(f"找到输出要素标记（分离匹配）: 索引 {i}, 文本: {repr(text)}")
            
            # 检查"不涉及"标记 - 增强检查逻辑（检查表格内和段落中）
            if found_input_text and input_marker_index == i:
                # 先检查当前段落是否包含"不涉及"
                if "不涉及" in text:
                    input_not_involved = True
                    parser_logger.debug(f"输入要素标记为不涉及（当前段落）: 索引 {i}")
                else:
                    # 检查后续段落（文档末尾需要检查更多行）
                    check_range = min(i + 20, search_end) if is_document_end else min(i + 10, search_end)
                    for j in range(i + 1, check_range):
                        next_para = self.paragraphs[j]
                        next_text = next_para.text.strip()
                        # 跳过空行
                        if not next_text:
                            continue
                        # 如果遇到下一个标记（如"输出要素"），停止检查
                        if any(variant in next_text for variant in output_marker_variants):
                            break
                        # 如果遇到下一个章节标记（如"三、"），停止检查
                        if re.match(r'^[一二三四五六七八九十]+、', next_text):
                            break
                        # 检查是否包含"不涉及"（支持表格内的情况）
                        if "不涉及" in next_text:
                            input_not_involved = True
                            parser_logger.debug(f"输入要素标记为不涉及（后续段落）: 索引 {j}")
                            break
            
            if found_output_text and output_marker_index == i:
                # 先检查当前段落是否包含"不涉及"
                if "不涉及" in text:
                    output_not_involved = True
                    parser_logger.debug(f"输出要素标记为不涉及（当前段落）: 索引 {i}")
                else:
                    # 检查后续段落（文档末尾需要检查更多行）
                    check_range = min(i + 20, search_end) if is_document_end else min(i + 10, search_end)
                    for j in range(i + 1, check_range):
                        next_para = self.paragraphs[j]
                        next_text = next_para.text.strip()
                        # 跳过空行
                        if not next_text:
                            continue
                        # 如果遇到下一个章节标记（如"三、"），停止检查
                        if re.match(r'^[一二三四五六七八九十]+、', next_text):
                            break
                        # 检查是否包含"不涉及"（支持表格内的情况）
                        if "不涉及" in next_text:
                            output_not_involved = True
                            parser_logger.debug(f"输出要素标记为不涉及（后续段落）: 索引 {j}")
                            break
        
        # 只有在不是"不涉及"的情况下才查找表格
        # 使用分级查找策略和回溯机制
        if found_input_text and not input_not_involved:
            parser_logger.info(f"开始查找输入要素表 - 标记索引: {input_marker_index}, 搜索结束: {search_end}")
            input_elements = self._find_table_after_marker_in_range(
                input_marker_index, search_end, is_input=True, is_document_end=is_document_end
            )
            # 文档末尾回退机制：如果没找到，尝试使用所有未使用的表格
            if not input_elements and is_document_end:
                parser_logger.warning(f"输入要素表未找到，尝试回退策略")
                input_elements = self._search_all_unused_tables(is_input=True)
        elif input_not_involved:
            # 如果标记为"不涉及"，不提取输入要素
            input_elements = []
            parser_logger.debug("输入要素标记为不涉及，跳过提取")
        
        if found_output_text and not output_not_involved:
            parser_logger.info(f"开始查找输出要素表 - 标记索引: {output_marker_index}, 搜索结束: {search_end}")
            output_elements = self._find_table_after_marker_in_range(
                output_marker_index, search_end, is_input=False, is_document_end=is_document_end
            )
            # 文档末尾回退机制：如果没找到，尝试使用所有未使用的表格
            if not output_elements and is_document_end:
                parser_logger.warning(f"输出要素表未找到，尝试回退策略")
                output_elements = self._search_all_unused_tables(is_input=False)
        elif output_not_involved:
            # 如果标记为"不涉及"，不提取输出要素
            output_elements = []
            parser_logger.debug("输出要素标记为不涉及，跳过提取")
        
        # 记录结果摘要
        parser_logger.info(
            f"提取完成 - 输入要素数量: {len(input_elements)}, 输出要素数量: {len(output_elements)}"
        )
        
        # 缓存结果
        result = (input_elements, output_elements)
        self._extraction_cache[cache_key] = result
        
        return result
    
    def _parse_input_table(self, table: Table) -> List[InputElement]:
        """解析输入要素表（增强版 - 增强容错能力）
        
        增强功能：
        1. 支持表头行可能不是第一行的情况
        2. 增强对空行和异常数据的处理
        3. 支持列数不一致的情况
        """
        elements = []
        
        if len(table.rows) < 2:
            parser_logger.warning("输入要素表行数不足，跳过解析")
            return elements
        
        # 获取表头 - 尝试第一行，如果第一行看起来不像表头，尝试第二行
        header_row_idx = 0
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        
        # 检查第一行是否像表头（包含"字段名称"等关键词）
        header_indicators = ["字段名称", "字段名", "名称", "是否必输", "类型"]
        first_row_looks_like_header = any(indicator in ' '.join(headers) for indicator in header_indicators)
        
        if not first_row_looks_like_header and len(table.rows) > 2:
            # 尝试第二行作为表头
            headers = [cell.text.strip() for cell in table.rows[1].cells]
            if any(indicator in ' '.join(headers) for indicator in header_indicators):
                header_row_idx = 1
                parser_logger.debug("使用第二行作为表头")
        
        # 使用模糊匹配查找各列索引（支持更多变体）
        index_idx = 0
        name_idx = self._fuzzy_find_column_index(headers, ["字段名称", "名称", "字段名", "栏位名称"])
        required_idx = self._fuzzy_find_column_index(headers, ["是否必输", "是否必填", "必输", "必填", "必填项"])
        type_idx = self._fuzzy_find_column_index(headers, ["类型", "字段类型", "数据类型"])
        precision_idx = self._fuzzy_find_column_index(headers, ["精度", "小数位数"])
        format_idx = self._fuzzy_find_column_index(headers, ["字段格式", "格式", "输入格式", "数据格式"])
        limit_idx = self._fuzzy_find_column_index(headers, ["输入限制", "限制", "数据字典", "取值范围"])
        desc_idx = self._fuzzy_find_column_index(headers, ["说明", "描述", "备注", "注释"])
        
        if name_idx == -1:
            parser_logger.warning(f"输入要素表未找到字段名称列，表头: {headers[:5]}")
            return elements
        
        parser_logger.debug(
            f"输入要素表列索引 - 名称: {name_idx}, 必输: {required_idx}, 类型: {type_idx}, "
            f"格式: {format_idx}, 限制: {limit_idx}, 说明: {desc_idx}"
        )
        
        # 解析数据行（从表头行之后开始）
        data_start_row = header_row_idx + 1
        for i, row in enumerate(table.rows[data_start_row:], start=1):
            # 容错：检查行是否有足够的单元格
            if len(row.cells) < name_idx + 1:
                parser_logger.debug(f"输入要素表第{data_start_row + i}行单元格数不足，跳过")
                continue
            
            try:
                cells = [cell.text.strip() for cell in row.cells]
            except Exception as e:
                parser_logger.warning(f"解析输入要素表第{data_start_row + i}行时出错: {e}")
                continue
            
            # 跳过空行（字段名称为空）
            if name_idx < len(cells) and not cells[name_idx]:
                continue
            
            # 容错：解析索引
            try:
                if index_idx < len(cells) and cells[index_idx]:
                    index = int(cells[index_idx])
                else:
                    index = i
            except (ValueError, IndexError):
                index = i
            
            # 容错：获取字段名称
            field_name = cells[name_idx] if name_idx < len(cells) else ""
            if not field_name:
                continue
            
            # 容错：获取其他字段（使用安全访问）
            def safe_get_cell(idx: int, default: str = "") -> str:
                if idx == -1 or idx >= len(cells):
                    return default
                return cells[idx] if cells[idx] else default
            
            element = InputElement(
                index=index,
                field_name=field_name,
                required=safe_get_cell(required_idx, "否"),
                field_type=safe_get_cell(type_idx) or None,
                precision=safe_get_cell(precision_idx) or None,
                field_format=safe_get_cell(format_idx) or None,
                input_limit=safe_get_cell(limit_idx) or None,
                description=safe_get_cell(desc_idx) or None
            )
            elements.append(element)
        
        parser_logger.debug(f"输入要素表解析完成，提取 {len(elements)} 个要素")
        return elements
    
    def _parse_output_table(self, table: Table) -> List[OutputElement]:
        """解析输出要素表（增强版 - 增强容错能力）
        
        增强功能：
        1. 支持表头行可能不是第一行的情况
        2. 增强对空行和异常数据的处理
        3. 支持列数不一致的情况
        """
        elements = []
        
        if len(table.rows) < 2:
            parser_logger.warning("输出要素表行数不足，跳过解析")
            return elements
        
        # 获取表头 - 尝试第一行，如果第一行看起来不像表头，尝试第二行
        header_row_idx = 0
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        
        # 检查第一行是否像表头（包含"字段名称"和"类型"等关键词）
        header_indicators = ["字段名称", "字段名", "名称", "类型"]
        first_row_looks_like_header = any(indicator in ' '.join(headers) for indicator in header_indicators)
        
        if not first_row_looks_like_header and len(table.rows) > 2:
            # 尝试第二行作为表头
            headers = [cell.text.strip() for cell in table.rows[1].cells]
            if any(indicator in ' '.join(headers) for indicator in header_indicators):
                header_row_idx = 1
                parser_logger.debug("使用第二行作为表头")
        
        # 使用模糊匹配查找各列索引（支持更多变体）
        index_idx = 0
        name_idx = self._fuzzy_find_column_index(headers, ["字段名称", "名称", "字段名", "栏位名称"])
        type_idx = self._fuzzy_find_column_index(headers, ["类型", "字段类型", "数据类型"])
        precision_idx = self._fuzzy_find_column_index(headers, ["精度", "小数位数"])
        format_idx = self._fuzzy_find_column_index(headers, ["字段格式", "格式", "输出格式", "数据格式"])
        desc_idx = self._fuzzy_find_column_index(headers, ["说明", "描述", "备注", "注释"])
        
        if name_idx == -1:
            parser_logger.warning(f"输出要素表未找到字段名称列，表头: {headers[:5]}")
            return elements
        
        # 输出表必须包含"类型"列
        if type_idx == -1:
            parser_logger.warning(f"输出要素表未找到类型列，表头: {headers[:5]}")
            # 不返回空，尝试继续解析（可能类型列名称不同）
        
        parser_logger.debug(
            f"输出要素表列索引 - 名称: {name_idx}, 类型: {type_idx}, "
            f"格式: {format_idx}, 说明: {desc_idx}"
        )
        
        # 解析数据行（从表头行之后开始）
        data_start_row = header_row_idx + 1
        for i, row in enumerate(table.rows[data_start_row:], start=1):
            # 容错：检查行是否有足够的单元格
            if len(row.cells) < name_idx + 1:
                parser_logger.debug(f"输出要素表第{data_start_row + i}行单元格数不足，跳过")
                continue
            
            try:
                cells = [cell.text.strip() for cell in row.cells]
            except Exception as e:
                parser_logger.warning(f"解析输出要素表第{data_start_row + i}行时出错: {e}")
                continue
            
            # 跳过空行（字段名称为空）
            if name_idx < len(cells) and not cells[name_idx]:
                continue
            
            # 容错：解析索引
            try:
                if index_idx < len(cells) and cells[index_idx]:
                    index = int(cells[index_idx])
                else:
                    index = i
            except (ValueError, IndexError):
                index = i
            
            # 容错：获取字段名称
            field_name = cells[name_idx] if name_idx < len(cells) else ""
            if not field_name:
                continue
            
            # 容错：获取其他字段（使用安全访问）
            def safe_get_cell(idx: int, default: str = "") -> str:
                if idx == -1 or idx >= len(cells):
                    return default
                return cells[idx] if cells[idx] else default
            
            element = OutputElement(
                index=index,
                field_name=field_name,
                field_type=safe_get_cell(type_idx) or None,
                precision=safe_get_cell(precision_idx) or None,
                field_format=safe_get_cell(format_idx) or None,
                description=safe_get_cell(desc_idx) or None
            )
            elements.append(element)
        
        parser_logger.debug(f"输出要素表解析完成，提取 {len(elements)} 个要素")
        return elements
    
    def _is_heading(self, para: Paragraph, level: int) -> bool:
        """判断段落是否是指定级别的标题"""
        style_name = para.style.name
        # 检查样式名称（Heading 1, Heading 2等）
        if style_name.startswith('Heading'):
            try:
                heading_level = int(style_name.replace('Heading ', ''))
                return heading_level == level
            except:
                return False
        return False
    
    def _find_column_index(self, headers: List[str], keywords: List[str]) -> int:
        """查找包含关键词的列索引"""
        for i, header in enumerate(headers):
            for keyword in keywords:
                if keyword in header:
                    return i
        return -1
    
    def _fuzzy_find_column_index(self, headers: List[str], keywords: List[str]) -> int:
        """模糊查找列索引
        - 先尝试完全匹配
        - 再尝试部分包含
        - 最后尝试相似匹配
        """
        # 1. 完全匹配
        for i, header in enumerate(headers):
            for keyword in keywords:
                if keyword == header or keyword in header:
                    return i
        
        # 2. 部分包含（去除空格和标点）
        for i, header in enumerate(headers):
            cleaned_header = re.sub(r'[^\w\u4e00-\u9fa5]', '', header)
            for keyword in keywords:
                cleaned_keyword = re.sub(r'[^\w\u4e00-\u9fa5]', '', keyword)
                if cleaned_keyword in cleaned_header or cleaned_header in cleaned_keyword:
                    return i
        
        return -1
    
    def _is_input_table(self, header_text: str, table: Table = None) -> bool:
        """判断是否为输入要素表（增强版 - 支持更灵活的表头识别和语义识别）
        
        Args:
            header_text: 表头文本
            table: 表格对象（可选，用于检查表体内容）
        """
        # 1. 基础检查：宽松匹配字段名称（支持多种变体）
        field_name_indicators = ["字段名称", "字段名", "名称", "栏位名称"]
        if not any(indicator in header_text for indicator in field_name_indicators):
            return False
        
        # 2. 排除文件格式表格：如果表头明确包含文件格式相关字段，不是输入要素表
        file_format_keywords = ["文件头begin", "文件体begin", "文件尾begin", "文件头end", "文件体end", "文件尾end", 
                               "文件总记录数量", "摘要代码分类代码", "文件头Begin", "文件体Begin", "文件尾Begin"]
        if any(keyword in header_text for keyword in file_format_keywords):
            return False
        
        # 3. 检查表体内容（如果提供了表格对象）- 增强语义识别
        if table and len(table.rows) > 1:
            # 检查前3行数据，提高识别准确性
            for row_idx in range(1, min(4, len(table.rows))):
                row_text = ' '.join([cell.text.strip() for cell in table.rows[row_idx].cells])
                # 如果数据行包含文件格式标记，排除
                if any(keyword in row_text for keyword in ["文件头begin", "文件体begin", "文件尾begin", 
                                                           "文件头Begin", "文件体Begin", "文件尾Begin"]):
                    return False
                # 如果数据行包含明显的输入表特征（如"是否必输"的值），增强判断
                if any(indicator in row_text for indicator in ["是", "否"]) and "类型" not in header_text:
                    # 可能是输入表（包含是否必输列）
                    pass
        
        # 4. 优先判断：输入表典型特征
        input_indicators = ["是否必输", "是否必填", "必输", "必填", "数据来源", "输入限制"]
        if any(indicator in header_text for indicator in input_indicators):
            # 进一步确认：排除输出表特征
            if "输出限制" in header_text:
                return False
            # 如果同时包含"类型"，需要检查是否有输入表特征
            if "类型" in header_text:
                # 如果同时有"是否必输"或"数据来源"，优先判断为输入表
                if any(ind in header_text for ind in ["是否必输", "是否必填", "数据来源"]):
                    return True
            else:
                return True
        
        # 5. 排除输出表特征
        # 如果包含"类型"但没有输入表特征，且没有"是否必输"和"数据来源"，可能是输出表
        if "类型" in header_text and "是否必输" not in header_text and "是否必填" not in header_text and "数据来源" not in header_text:
            # 检查是否有输出表特征
            if "输出限制" in header_text or ("输出" in header_text and "类型" in header_text):
                return False
        
        # 6. 如果只有字段名称列，没有其他明显特征，默认不判断为输入表（需要更多上下文）
        # 这种情况下，依赖调用方传入的is_input参数和表格位置来判断
        return False
    
    def _is_output_table(self, header_text: str, table: Table = None) -> bool:
        """判断是否为输出要素表（增强版 - 支持更灵活的表头识别和语义识别）
        
        Args:
            header_text: 表头文本
            table: 表格对象（可选，用于检查表体内容）
        """
        # 1. 基础检查：宽松匹配字段名称（支持多种变体）
        field_name_indicators = ["字段名称", "字段名", "名称", "栏位名称"]
        if not any(indicator in header_text for indicator in field_name_indicators):
            return False
        
        # 2. 排除文件格式表格：如果表头明确包含文件格式相关字段，不是输出要素表
        file_format_keywords = ["文件头begin", "文件体begin", "文件尾begin", "文件头end", "文件体end", "文件尾end", 
                               "文件总记录数量", "摘要代码分类代码", "文件头Begin", "文件体Begin", "文件尾Begin"]
        if any(keyword in header_text for keyword in file_format_keywords):
            return False
        
        # 3. 检查表体内容（如果提供了表格对象）- 增强语义识别
        if table and len(table.rows) > 1:
            # 检查前3行数据，提高识别准确性
            for row_idx in range(1, min(4, len(table.rows))):
                row_text = ' '.join([cell.text.strip() for cell in table.rows[row_idx].cells])
                # 如果数据行包含文件格式标记，排除
                if any(keyword in row_text for keyword in ["文件头begin", "文件体begin", "文件尾begin",
                                                           "文件头Begin", "文件体Begin", "文件尾Begin"]):
                    return False
        
        # 4. 必须包含"类型"（输出表的典型特征）- 但允许其他列名变体
        type_indicators = ["类型", "字段类型", "数据类型"]
        if not any(indicator in header_text for indicator in type_indicators):
            return False
        
        # 5. 排除输入表特征：如果包含"是否必输"、"是否必填"或"数据来源"，不是输出表
        if any(ind in header_text for ind in ["是否必输", "是否必填", "必输", "必填", "数据来源"]):
            return False
        
        # 6. 包含"类型"且不包含"是否必输"和"数据来源"，是输出表
        # 注意：即使有"输入限制"列，只要没有"是否必输"和"数据来源"，也是输出表
        return True
    
    def _search_tables_near_marker(self, marker_index: int, is_input: bool, max_distance: int = 20) -> List:
        """在标记附近搜索表格"""
        elements = []
        start_idx = max(0, marker_index - max_distance)
        end_idx = min(len(self.paragraphs), marker_index + max_distance)
        
        # 在范围内查找表格（通过检查段落和表格的关联）
        # 由于python-docx无法直接关联段落和表格，我们采用顺序查找策略
        # 找到标记后，按顺序查找后续的未使用表格
        for table_idx, table in enumerate(self.tables):
            if table_idx in self.used_tables:
                continue
            if len(table.rows) < 2:
                continue
            
            header_text = ' '.join([cell.text.strip() for cell in table.rows[0].cells])
            
            if is_input and self._is_input_table(header_text, table):
                parsed = self._parse_input_table(table)
                if parsed:
                    elements = parsed
                    self.used_tables.add(table_idx)
                    break
            elif not is_input and self._is_output_table(header_text, table):
                parsed = self._parse_output_table(table)
                if parsed:
                    elements = parsed
                    self.used_tables.add(table_idx)
                    break
        
        return elements
    
    def _search_tables_in_range(self, start: int, end: int, is_input: bool, allow_used: bool = False) -> List:
        """在指定范围内搜索表格"""
        elements = []
        
        # 按顺序查找表格
        for table_idx, table in enumerate(self.tables):
            if not allow_used and table_idx in self.used_tables:
                continue
            if len(table.rows) < 2:
                continue
            
            header_text = ' '.join([cell.text.strip() for cell in table.rows[0].cells])
            
            if is_input and self._is_input_table(header_text):
                parsed = self._parse_input_table(table)
                if parsed:
                    elements = parsed
                    if not allow_used:
                        self.used_tables.add(table_idx)
                    break
            elif not is_input and self._is_output_table(header_text):
                parsed = self._parse_output_table(table)
                if parsed:
                    elements = parsed
                    if not allow_used:
                        self.used_tables.add(table_idx)
                    break
        
        return elements
    
    def _search_all_unused_tables(self, is_input: bool) -> List:
        """遍历所有未使用的表格"""
        elements = []
        
        for table_idx, table in enumerate(self.tables):
            if table_idx in self.used_tables:
                continue
            if len(table.rows) < 2:
                continue
            
            header_text = ' '.join([cell.text.strip() for cell in table.rows[0].cells])
            
            if is_input and self._is_input_table(header_text, table):
                parsed = self._parse_input_table(table)
                if parsed:
                    elements = parsed
                    self.used_tables.add(table_idx)
                    break
            elif not is_input and self._is_output_table(header_text, table):
                parsed = self._parse_output_table(table)
                if parsed:
                    elements = parsed
                    self.used_tables.add(table_idx)
                    break
        
        return elements
    
    def _find_table_after_marker_in_range(self, marker_index: int, end_index: int, is_input: bool, is_document_end: bool = False, function_name: Optional[str] = None) -> List:
        """在标记后的指定范围内查找表格（优化版 - 分级查找策略和绝对位置计算）
        
        策略（分级查找）：
        1. 首选：标记后最近的未使用表格（使用绝对位置计算）
        2. 次选：标记后最近的所有表格（忽略已使用标记）
        3. 最后：按内容语义匹配的表格
        
        Args:
            marker_index: 标记的段落索引
            end_index: 搜索结束的段落索引
            is_input: 是否为输入表
            is_document_end: 是否为文档末尾步骤
        """
        elements = []
        table_type = "输入" if is_input else "输出"
        
        parser_logger.debug(
            f"开始查找{table_type}要素表 - 标记索引: {marker_index}, 结束索引: {end_index}, "
            f"文档末尾: {is_document_end}"
        )
        
        # 文档末尾：使用绝对位置计算，估算表格在文档中的位置
        # 由于python-docx无法直接获取表格的段落位置，我们使用表格索引来估算
        # 假设表格在文档中大致均匀分布
        doc_length = len(self.paragraphs)
        total_tables = len(self.tables)
        
        # 计算标记的绝对位置（段落索引）
        marker_abs_pos = marker_index
        
        # 估算每个表格对应的段落位置范围
        # 这是一个近似值，但比相对位置更准确
        def estimate_table_position(table_idx: int) -> int:
            """估算表格在文档中的段落位置"""
            if total_tables == 0:
                return 0
            # 假设表格均匀分布，但允许一定误差
            estimated_pos = int((table_idx / total_tables) * doc_length)
            return estimated_pos
        
        # 策略1：首选 - 标记后最近的未使用表格（优先在章节范围内）
        parser_logger.debug(f"策略1: 查找标记后最近的未使用{table_type}要素表")
        candidate_tables_level1 = []  # (table_idx, parsed_elements, estimated_pos, header_summary)
        candidate_tables_level1_in_range = []  # 在章节范围内的表格
        
        for table_idx, table in enumerate(self.tables):
            # 跳过已使用的表格
            if table_idx in self.used_tables:
                continue
            
            # 跳过行数不足的表格
            if len(table.rows) < 2:
                continue
            
            # 估算表格位置
            table_estimated_pos = estimate_table_position(table_idx)
            
            # 检查表格是否在章节范围内（优先使用章节范围判断）
            # 对于非文档末尾的情况，如果表格在章节范围内，即使估算位置在标记前也接受（因为估算可能不准确）
            if is_document_end:
                # 文档末尾：允许表格位置在标记前10%范围内（因为估算不准确）
                if table_estimated_pos < marker_abs_pos - doc_length * 0.1:
                    continue
                is_in_range = marker_abs_pos <= table_estimated_pos <= end_index
            else:
                # 非文档末尾：优先使用章节范围判断
                # 如果表格在章节范围内（table_estimated_pos <= end_index），即使估算位置在标记前也接受
                # 但需要确保表格在功能章节内（从功能开始位置到章节结束位置）
                # 这里我们使用更宽松的判断：只要表格在章节结束位置之前，就认为在范围内
                # 因为估算位置可能不准确，表格实际位置可能在标记附近
                is_in_range = table_estimated_pos <= end_index
                # 但如果估算位置明显在标记前太多（超过功能章节开始位置），则排除（可能是前面的功能）
                # 使用功能章节开始位置作为参考（如果有的话），否则使用标记位置
                # 这里我们放宽限制，只要表格在章节结束位置之前就接受
                # 但如果估算位置明显在标记前太多（超过50%文档长度），则排除
                if table_estimated_pos < marker_abs_pos - doc_length * 0.5:
                    continue
            
            header_text = ' '.join([cell.text.strip() for cell in table.rows[0].cells])
            header_summary = header_text[:50]  # 保留表头摘要用于日志
            
            # 检查数据行，排除文件格式表格
            is_file_format = False
            if len(table.rows) > 1:
                # 检查前3行数据
                for row_idx in range(1, min(4, len(table.rows))):
                    row_text = ' '.join([cell.text.strip() for cell in table.rows[row_idx].cells])
                    file_format_keywords = ["文件头begin", "文件体begin", "文件尾begin", "文件头end", 
                                           "文件体end", "文件尾end", "文件头Begin", "文件体Begin", "文件尾Begin"]
                    if any(keyword in row_text for keyword in file_format_keywords):
                        is_file_format = True
                        break
            
            if is_file_format:
                parser_logger.debug(f"跳过文件格式表格: 索引 {table_idx}")
                continue
            
            # 使用精确的表格类型判断（传入表格对象以检查表体内容）
            if is_input and self._is_input_table(header_text, table):
                parsed = self._parse_input_table(table)
                if parsed:
                    candidate = (table_idx, parsed, table_estimated_pos, header_summary)
                    if is_in_range:
                        candidate_tables_level1_in_range.append(candidate)
                        parser_logger.debug(
                            f"找到候选输入表（策略1，章节内）: 索引 {table_idx}, 估算位置 {table_estimated_pos}, "
                            f"要素数量 {len(parsed)}, 表头: {header_summary}"
                        )
                    else:
                        candidate_tables_level1.append(candidate)
                        parser_logger.debug(
                            f"找到候选输入表（策略1，章节外）: 索引 {table_idx}, 估算位置 {table_estimated_pos}, "
                            f"要素数量 {len(parsed)}, 表头: {header_summary}"
                        )
            elif not is_input and self._is_output_table(header_text, table):
                parsed = self._parse_output_table(table)
                if parsed:
                    candidate = (table_idx, parsed, table_estimated_pos, header_summary)
                    if is_in_range:
                        candidate_tables_level1_in_range.append(candidate)
                        parser_logger.debug(
                            f"找到候选输出表（策略1，章节内）: 索引 {table_idx}, 估算位置 {table_estimated_pos}, "
                            f"要素数量 {len(parsed)}, 表头: {header_summary}"
                        )
                    else:
                        candidate_tables_level1.append(candidate)
                        parser_logger.debug(
                            f"找到候选输出表（策略1，章节外）: 索引 {table_idx}, 估算位置 {table_estimated_pos}, "
                            f"要素数量 {len(parsed)}, 表头: {header_summary}"
                        )
        
        # 优先选择在章节范围内的表格
        if candidate_tables_level1_in_range:
            # 优先选择在标记后的表格，如果都在标记前，则选择最接近标记的
            # 将表格分为两类：标记后的和标记前的
            tables_after_marker = [c for c in candidate_tables_level1_in_range if c[2] >= marker_abs_pos]
            tables_before_marker = [c for c in candidate_tables_level1_in_range if c[2] < marker_abs_pos]
            
            if tables_after_marker:
                # 优先选择标记后的表格，按距离标记位置排序
                tables_after_marker.sort(key=lambda x: x[2] - marker_abs_pos)  # 按距离标记的距离排序
                table_idx, parsed, estimated_pos, header_summary = tables_after_marker[0]
            else:
                # 如果都在标记前，选择最接近标记的
                tables_before_marker.sort(key=lambda x: marker_abs_pos - x[2])  # 按距离标记的距离排序
                table_idx, parsed, estimated_pos, header_summary = tables_before_marker[0]
            
            elements = parsed
            self.used_tables.add(table_idx)
            parser_logger.info(
                f"策略1成功（章节范围内） - 找到{table_type}要素表: 索引 {table_idx}, 估算位置 {estimated_pos}, "
                f"要素数量 {len(parsed)}, 表头: {header_summary}"
            )
            return elements
        
        # 如果章节范围内没有，再选择章节范围外的
        if candidate_tables_level1:
            # 按估算位置排序，选择最接近标记的表格
            candidate_tables_level1.sort(key=lambda x: abs(x[2] - marker_abs_pos))
            table_idx, parsed, estimated_pos, header_summary = candidate_tables_level1[0]
            elements = parsed
            self.used_tables.add(table_idx)
            parser_logger.info(
                f"策略1成功（章节范围外） - 找到{table_type}要素表: 索引 {table_idx}, 估算位置 {estimated_pos}, "
                f"要素数量 {len(parsed)}, 表头: {header_summary}"
            )
            return elements
        
        # 策略2：次选 - 标记后最近的所有表格（忽略已使用标记）
        # 优先选择在章节范围内的表格
        parser_logger.debug(f"策略1失败，尝试策略2: 查找标记后最近的所有{table_type}要素表（忽略已使用标记）")
        candidate_tables_level2 = []
        candidate_tables_level2_in_range = []  # 在章节范围内的表格
        
        for table_idx, table in enumerate(self.tables):
            # 跳过行数不足的表格
            if len(table.rows) < 2:
                continue
            
            # 估算表格位置
            table_estimated_pos = estimate_table_position(table_idx)
            
            # 只考虑标记后的表格
            if is_document_end:
                if table_estimated_pos < marker_abs_pos - doc_length * 0.1:
                    continue
            else:
                if table_estimated_pos < marker_abs_pos:
                    continue
            
            # 检查表格是否在章节范围内
            is_in_range = marker_abs_pos <= table_estimated_pos <= end_index
            
            header_text = ' '.join([cell.text.strip() for cell in table.rows[0].cells])
            header_summary = header_text[:50]
            
            # 检查文件格式表格
            is_file_format = False
            if len(table.rows) > 1:
                for row_idx in range(1, min(4, len(table.rows))):
                    row_text = ' '.join([cell.text.strip() for cell in table.rows[row_idx].cells])
                    file_format_keywords = ["文件头begin", "文件体begin", "文件尾begin", "文件头end", 
                                           "文件体end", "文件尾end", "文件头Begin", "文件体Begin", "文件尾Begin"]
                    if any(keyword in row_text for keyword in file_format_keywords):
                        is_file_format = True
                        break
            
            if is_file_format:
                continue
            
            # 表格类型判断（忽略已使用标记）
            if is_input and self._is_input_table(header_text, table):
                parsed = self._parse_input_table(table)
                if parsed:
                    candidate = (table_idx, parsed, table_estimated_pos, header_summary)
                    if is_in_range:
                        candidate_tables_level2_in_range.append(candidate)
                    else:
                        candidate_tables_level2.append(candidate)
            elif not is_input and self._is_output_table(header_text, table):
                parsed = self._parse_output_table(table)
                if parsed:
                    candidate = (table_idx, parsed, table_estimated_pos, header_summary)
                    if is_in_range:
                        candidate_tables_level2_in_range.append(candidate)
                    else:
                        candidate_tables_level2.append(candidate)
        
        # 优先选择在章节范围内的表格
        if candidate_tables_level2_in_range:
            candidate_tables_level2_in_range.sort(key=lambda x: abs(x[2] - marker_abs_pos))
            table_idx, parsed, estimated_pos, header_summary = candidate_tables_level2_in_range[0]
            elements = parsed
            self.used_tables.add(table_idx)
            parser_logger.info(
                f"策略2成功（章节范围内） - 找到{table_type}要素表: 索引 {table_idx}, 估算位置 {estimated_pos}, "
                f"要素数量 {len(parsed)}, 表头: {header_summary}"
            )
            return elements
        
        # 如果章节范围内没有，再选择章节范围外的
        if candidate_tables_level2:
            candidate_tables_level2.sort(key=lambda x: abs(x[2] - marker_abs_pos))
            table_idx, parsed, estimated_pos, header_summary = candidate_tables_level2[0]
            elements = parsed
            # 标记为已使用，避免其他功能找到同一个表格
            self.used_tables.add(table_idx)
            parser_logger.info(
                f"策略2成功（章节范围外） - 找到{table_type}要素表: 索引 {table_idx}, 估算位置 {estimated_pos}, "
                f"要素数量 {len(parsed)}, 表头: {header_summary}"
            )
            return elements
        
        # 策略3：最后 - 按内容语义匹配的表格（在整个文档中查找）
        # 优先选择在章节范围内的表格
        parser_logger.debug(f"策略2失败，尝试策略3: 按内容语义匹配查找{table_type}要素表")
        candidate_tables_level3 = []
        candidate_tables_level3_in_range = []  # 在章节范围内的表格
        
        for table_idx, table in enumerate(self.tables):
            # 跳过行数不足的表格
            if len(table.rows) < 2:
                continue
            
            header_text = ' '.join([cell.text.strip() for cell in table.rows[0].cells])
            header_summary = header_text[:50]
            
            # 检查文件格式表格
            is_file_format = False
            if len(table.rows) > 1:
                for row_idx in range(1, min(4, len(table.rows))):
                    row_text = ' '.join([cell.text.strip() for cell in table.rows[row_idx].cells])
                    file_format_keywords = ["文件头begin", "文件体begin", "文件尾begin", "文件头end", 
                                           "文件体end", "文件尾end", "文件头Begin", "文件体Begin", "文件尾Begin"]
                    if any(keyword in row_text for keyword in file_format_keywords):
                        is_file_format = True
                        break
            
            if is_file_format:
                continue
            
            # 估算表格位置
            table_estimated_pos = estimate_table_position(table_idx)
            # 检查表格是否在章节范围内
            is_in_range = marker_abs_pos <= table_estimated_pos <= end_index
            
            # 表格类型判断
            if is_input and self._is_input_table(header_text, table):
                parsed = self._parse_input_table(table)
                if parsed:
                    candidate = (table_idx, parsed, table_estimated_pos, header_summary)
                    if is_in_range:
                        candidate_tables_level3_in_range.append(candidate)
                    else:
                        candidate_tables_level3.append(candidate)
            elif not is_input and self._is_output_table(header_text, table):
                parsed = self._parse_output_table(table)
                if parsed:
                    candidate = (table_idx, parsed, table_estimated_pos, header_summary)
                    if is_in_range:
                        candidate_tables_level3_in_range.append(candidate)
                    else:
                        candidate_tables_level3.append(candidate)
        
        # 优先选择在章节范围内的表格
        if candidate_tables_level3_in_range:
            candidate_tables_level3_in_range.sort(key=lambda x: abs(x[2] - marker_abs_pos))
            table_idx, parsed, estimated_pos, header_summary = candidate_tables_level3_in_range[0]
            elements = parsed
            self.used_tables.add(table_idx)
            parser_logger.info(
                f"策略3成功（章节范围内） - 找到{table_type}要素表: 索引 {table_idx}, 估算位置 {estimated_pos}, "
                f"要素数量 {len(parsed)}, 表头: {header_summary}"
            )
            return elements
        
        # 如果章节范围内没有，再选择章节范围外的
        if candidate_tables_level3:
            candidate_tables_level3.sort(key=lambda x: abs(x[2] - marker_abs_pos))
            table_idx, parsed, estimated_pos, header_summary = candidate_tables_level3[0]
            elements = parsed
            self.used_tables.add(table_idx)
            parser_logger.info(
                f"策略3成功（章节范围外） - 找到{table_type}要素表: 索引 {table_idx}, 估算位置 {estimated_pos}, "
                f"要素数量 {len(parsed)}, 表头: {header_summary}"
            )
            return elements
        
        parser_logger.warning(f"所有策略均失败 - 未找到{table_type}要素表")
        return elements
    
    def _find_table_in_section_range(self, marker_index: int, end_index: int, is_input: bool, section_start: int) -> List:
        """在指定章节范围内查找表格
        
        策略：
        1. 在章节范围内（section_start到end_index），查找标记后最近的未使用表格
        2. 这样可以确保每个步骤匹配到该步骤自己的表格，而不是被其他步骤占用
        3. 排除文件格式表格（包含"文件头begin"等字段）
        """
        elements = []
        
        # 计算标记在文档中的相对位置
        marker_doc_pos = marker_index / len(self.paragraphs) if len(self.paragraphs) > 0 else 0
        
        # 收集所有候选表格
        candidate_tables = []
        for table_idx, table in enumerate(self.tables):
            if table_idx in self.used_tables:
                continue
            if len(table.rows) < 2:
                continue
            
            header_text = ' '.join([cell.text.strip() for cell in table.rows[0].cells])
            
            # 检查数据行，排除文件格式表格
            if len(table.rows) > 1:
                first_data_row = table.rows[1]
                first_data_text = ' '.join([cell.text.strip() for cell in first_data_row.cells])
                file_format_keywords = ["文件头begin", "文件体begin", "文件尾begin", "文件头end", "文件体end", "文件尾end"]
                if any(keyword in first_data_text for keyword in file_format_keywords):
                    continue
            
            # 使用精确的表格类型判断
            if is_input and self._is_input_table(header_text, table):
                parsed = self._parse_input_table(table)
                if parsed:
                    # 计算表格的相对位置（估算表格在文档中的位置）
                    table_relative_pos = table_idx / len(self.tables) if len(self.tables) > 0 else 0
                    # 只考虑位置在标记后面的表格
                    if table_relative_pos >= marker_doc_pos - 0.1:
                        candidate_tables.append((table_idx, parsed, table_relative_pos))
            elif not is_input and self._is_output_table(header_text, table):
                parsed = self._parse_output_table(table)
                if parsed:
                    table_relative_pos = table_idx / len(self.tables) if len(self.tables) > 0 else 0
                    if table_relative_pos >= marker_doc_pos - 0.1:
                        candidate_tables.append((table_idx, parsed, table_relative_pos))
        
        # 如果有候选表格，优先选择位置在标记后面的表格
        if candidate_tables:
            # 优先选择位置在标记后面的表格
            after_marker_tables = [t for t in candidate_tables if t[2] >= marker_doc_pos - 0.05]
            if after_marker_tables:
                # 在标记后面的表格中，选择索引最小的（第一个匹配的）
                after_marker_tables.sort(key=lambda x: x[0])
                table_idx, parsed, _ = after_marker_tables[0]
            else:
                # 如果没有在标记后面的表格，选择索引最小的
                candidate_tables.sort(key=lambda x: x[0])
                table_idx, parsed, _ = candidate_tables[0]
            
            elements = parsed
            self.used_tables.add(table_idx)
            return elements
        
        return elements
    
    def _find_nearest_table_after_marker(self, marker_index: int, is_input: bool) -> List:
        """查找标记后最近的表格
        
        策略：
        1. 先查找标记后未使用的匹配表格（优先选择索引较大的，更可能在标记后面）
        2. 如果没找到，再查找所有未使用的匹配表格
        3. 使用精确的表格类型判断（_is_input_table和_is_output_table）
        """
        elements = []
        
        # 策略1：查找未使用的匹配表格，优先选择索引较大的（更可能在标记后面）
        candidate_tables = []
        for table_idx, table in enumerate(self.tables):
            if table_idx in self.used_tables:
                continue
            if len(table.rows) < 2:
                continue
            
            header_text = ' '.join([cell.text.strip() for cell in table.rows[0].cells])
            
            if is_input and self._is_input_table(header_text):
                parsed = self._parse_input_table(table)
                if parsed:
                    candidate_tables.append((table_idx, parsed))
            elif not is_input and self._is_output_table(header_text):
                parsed = self._parse_output_table(table)
                if parsed:
                    candidate_tables.append((table_idx, parsed))
        
        # 如果有多个候选表格，优先选择索引较大的（更可能在标记后面）
        if candidate_tables:
            # 按表格索引排序，选择最大的（最接近标记位置的）
            candidate_tables.sort(key=lambda x: x[0], reverse=True)
            table_idx, parsed = candidate_tables[0]
            elements = parsed
            self.used_tables.add(table_idx)
            return elements
        
        # 策略2：如果没找到未使用的表格，查找所有表格（包括已使用的）
        # 但只返回第一个匹配的，避免重复匹配
        for table_idx, table in enumerate(self.tables):
            if len(table.rows) < 2:
                continue
            
            header_text = ' '.join([cell.text.strip() for cell in table.rows[0].cells])
            
            if is_input and self._is_input_table(header_text):
                parsed = self._parse_input_table(table)
                if parsed:
                    elements = parsed
                    if table_idx not in self.used_tables:
                        self.used_tables.add(table_idx)
                    return elements
            elif not is_input and self._is_output_table(header_text):
                parsed = self._parse_output_table(table)
                if parsed:
                    elements = parsed
                    if table_idx not in self.used_tables:
                        self.used_tables.add(table_idx)
                    return elements
        
        return elements
    
    # ========== 非建模需求解析方法 ==========
    
    def _extract_file_controlled_info(self) -> Tuple[Optional[str], Optional[str]]:
        """从文件受控信息表或文档受控信息表提取文件编号和文件名称"""
        file_number = None
        file_name = None
        
        # 查找包含"文件受控信息"或"文档受控信息"的段落或表格
        for para in self.paragraphs[:100]:
            text = para.text.strip()
            if "文件受控信息" in text or "文档受控信息" in text:
                # 查找后续的表格
                for table in self.tables:
                    if len(table.rows) < 2:
                        continue
                    
                    header_row = table.rows[0]
                    headers = [cell.text.strip() for cell in header_row.cells]
                    header_text = ' '.join(headers)
                    
                    # 检查表格是否包含"文档受控信息"（可能是表头就是"文档受控信息"）
                    if "文档受控信息" in header_text:
                        # 特殊格式：行1可能是 ['文件编号', '值', '文件名称', '值']
                        if len(table.rows) >= 2:
                            data_row = table.rows[1]
                            cells = [cell.text.strip() for cell in data_row.cells]
                            
                            # 查找"文件编号"和"文件名称"的位置
                            for i, cell_text in enumerate(cells):
                                if "文件编号" in cell_text and i + 1 < len(cells):
                                    # 下一个单元格是文件编号的值
                                    file_number = cells[i + 1] if cells[i + 1] and cells[i + 1] != '/' else None
                                elif "文件名称" in cell_text and i + 1 < len(cells):
                                    # 下一个单元格是文件名称的值
                                    file_name = cells[i + 1] if cells[i + 1] and cells[i + 1] != '/' else None
                            
                            if file_number or file_name:
                                return file_number, file_name
                        
                        # 也尝试纵向布局：第一列是键，第二列是值
                        for row in table.rows[1:]:
                            if len(row.cells) >= 2:
                                key = row.cells[0].text.strip()
                                value = row.cells[1].text.strip()
                                
                                if value and value != '/':
                                    if ("文件编号" in key or ("编号" in key and "文件" in key)) and not file_number:
                                        file_number = value
                                    elif ("文件名称" in key or ("名称" in key and "文件" in key)) and not file_name:
                                        file_name = value
                        
                        if file_number or file_name:
                            return file_number, file_name
                    
                    # 查找文件编号和文件名称列（标准表格格式）
                    file_number_idx = self._find_column_index(headers, ["文件编号"])
                    file_name_idx = self._find_column_index(headers, ["文件名称"])
                    
                    if file_number_idx >= 0 or file_name_idx >= 0:
                        # 解析数据行（可能是横向布局：第一行是表头，第二行是值）
                        if len(table.rows) >= 2:
                            value_row = table.rows[1]
                            values = [cell.text.strip() for cell in value_row.cells]
                            
                            if file_number_idx >= 0 and file_number_idx < len(values):
                                file_number = values[file_number_idx] if values[file_number_idx] and values[file_number_idx] != '/' else None
                            
                            if file_name_idx >= 0 and file_name_idx < len(values):
                                file_name = values[file_name_idx] if values[file_name_idx] and values[file_name_idx] != '/' else None
                        
                        # 也尝试纵向布局：第一列是键，第二列是值
                        for row in table.rows[1:]:
                            if len(row.cells) >= 2:
                                key = row.cells[0].text.strip()
                                value = row.cells[1].text.strip()
                                
                                if value and value != '/':
                                    if "文件编号" in key and not file_number:
                                        file_number = value
                                    elif "文件名称" in key and not file_name:
                                        file_name = value
                        
                        if file_number or file_name:
                            return file_number, file_name
        
        # 也尝试直接从表格中查找（不依赖段落文本）
        for table in self.tables:
            if len(table.rows) < 2:
                continue
            
            header_row = table.rows[0]
            headers = [cell.text.strip() for cell in header_row.cells]
            header_text = ' '.join(headers)
            
            # 检查是否是文档受控信息表
            if "文档受控信息" in header_text:
                # 特殊格式处理
                if len(table.rows) >= 2:
                    data_row = table.rows[1]
                    cells = [cell.text.strip() for cell in data_row.cells]
                    
                    for i, cell_text in enumerate(cells):
                        if "文件编号" in cell_text and i + 1 < len(cells):
                            file_number = cells[i + 1] if cells[i + 1] and cells[i + 1] != '/' else None
                        elif "文件名称" in cell_text and i + 1 < len(cells):
                            file_name = cells[i + 1] if cells[i + 1] and cells[i + 1] != '/' else None
                
                # 也尝试纵向布局
                for row in table.rows[1:]:
                    if len(row.cells) >= 2:
                        key = row.cells[0].text.strip()
                        value = row.cells[1].text.strip()
                        
                        if value and value != '/':
                            if ("文件编号" in key or ("编号" in key and "文件" in key)) and not file_number:
                                file_number = value
                            elif ("文件名称" in key or ("名称" in key and "文件" in key)) and not file_name:
                                file_name = value
                
                if file_number or file_name:
                    return file_number, file_name
        
        return file_number, file_name
    
    def _extract_requirement_name(self, file_name: Optional[str]) -> Optional[str]:
        """提取需求名称
        方案一（主）：从文件名称中提取核心功能名
        方案二（备）：从功能清单第一项提取
        """
        # 方案一：从文件名称提取
        if file_name:
            # 清理换行符和多余空格
            file_name = file_name.replace('\n', '').replace('\r', '')
            file_name = re.sub(r'\s+', '', file_name)
            
            # 尝试多种正则模式匹配
            patterns = [
                r"大信贷系统(.+?)业务需求说明书",  # 标准格式
                r"大信贷系统详细业务-(.+?)需求说明书",  # 详细业务格式
                r"大信贷系统(.+?)需求说明书",  # 简化格式
            ]
            
            for pattern in patterns:
                match = re.search(pattern, file_name)
                if match:
                    requirement_name = match.group(1).strip()
                    # 清理处理：去除括号内容（但保留功能名称中的括号，如"贷款当日冲正（前台）"）
                    # 只去除说明性的括号，如"（优化对客服务类）"
                    # 如果需求名称本身包含括号，应该保留
                    if requirement_name:
                        return requirement_name
        
        # 方案二：从功能清单第一项提取
        functions = self._extract_function_list()
        if functions:
            return functions[0]
        
        return None
    
    def _extract_designer(self) -> Optional[str]:
        """提取设计者（作者）
        
        注意：设计者字段用于填写测试人员，不需要从文档中提取，直接返回None
        """
        return None
    
    def _extract_function_list(self) -> List[str]:
        """提取功能清单（仅功能名称列表）"""
        functions = []
        
        # 查找"5 功能*（A阶段）"章节，忽略之前的内容（非建模需求）
        function_chapter_start = -1
        for i, para in enumerate(self.paragraphs):
            text = para.text.strip()
            if "功能" in text and "（A阶段）" in text and re.match(r'^5\s+功能', text):
                function_chapter_start = i
                break
        
        # 查找"5.1 功能清单"章节（在功能章节之后）
        search_start = function_chapter_start if function_chapter_start >= 0 else 0
        for i in range(search_start, len(self.paragraphs)):
            text = self.paragraphs[i].text.strip()
            
            # 查找"5 功能*（A阶段）"或"5.1 功能清单"
            if ("功能" in text and "（A阶段）" in text) or "功能清单" in text:
                # 查找后续的表格
                for table in self.tables:
                    if len(table.rows) < 2:
                        continue
                    
                    header_row = table.rows[0]
                    headers = [cell.text.strip() for cell in header_row.cells]
                    
                    # 查找"业务功能名称"列
                    function_name_idx = self._find_column_index(headers, ["业务功能名称", "功能名称"])
                    
                    if function_name_idx >= 0:
                        # 解析数据行
                        for row in table.rows[1:]:
                            if len(row.cells) > function_name_idx:
                                function_name = row.cells[function_name_idx].text.strip()
                                if function_name and function_name not in functions:
                                    functions.append(function_name)
                        
                        if functions:
                            return functions
        
        return functions
    
    def _extract_function_titles_from_section(self, function_section_start: int) -> List[tuple]:
        """从功能说明部分提取功能标题（如 5.2.1、5.2.2 等）
        
        Returns:
            List[tuple]: [(function_name, index), ...] 功能名称和段落索引的列表
        """
        function_titles = []
        # 从功能说明部分开始搜索，不要跳过（因为功能标题可能在功能说明部分之后立即出现）
        # 但跳过前100个段落是为了跳过目录，如果功能说明部分在前100个段落之后，则从功能说明部分开始
        search_start = function_section_start + 1  # 从功能说明部分的下一个段落开始
        search_end = len(self.paragraphs)
        
        # 限制搜索范围：在功能说明部分后的200个段落内查找功能标题
        # 这样可以避免搜索到其他章节的功能
        search_end = min(search_start + 200, search_end)
        
        parser_logger.debug(f"在功能说明部分提取功能标题，搜索范围: {search_start} 到 {search_end}")
        
        for i in range(search_start, search_end):
            para = self.paragraphs[i]
            text = para.text.strip()
            
            if not text:
                continue
            
            # 匹配功能标题格式：5.2.1、5.2.2 等，后面跟着功能名称
            # 格式：5.2.1 查询账单明细功能 或 5.2.2 查询账单明细（催收）功能
            # 使用更宽松的匹配：允许数字编号后跟任意空白字符，然后提取后面的内容
            match = re.match(r'^(\d+\.\d+\.\d+)[\s\t]+(.+)', text)
            if match:
                function_number = match.group(1)  # 如 5.2.1
                function_title = match.group(2).strip()  # 如 查询账单明细功能 或 查询账单明细（催收）功能
                
                # 清理功能标题：移除页码（如末尾的数字）和多余的制表符/空格
                # 页码通常在末尾，格式如 "功能名称\t23" 或 "功能名称 23"
                function_title = re.sub(r'[\s\t]+\d+$', '', function_title)  # 移除末尾的制表符/空格+数字
                function_title = re.sub(r'\t+', ' ', function_title)  # 将制表符替换为空格
                function_title = function_title.strip()
                
                # 验证提取的内容是否像功能名称（包含"功能"或包含"查询"、"账单"等关键词）
                if "功能" in function_title or any(kw in function_title for kw in ["查询", "账单", "明细", "管理", "处理", "期供"]):
                    # 不自动添加"功能"后缀，直接使用标题中的名称
                    # 例如："5.2.1 查询贷款期供明细功能\t23" -> "查询贷款期供明细功能"
                    
                    function_titles.append((function_title, i))
                    parser_logger.info(f"从功能说明部分提取功能标题: {function_title}, 索引: {i}, 原始文本: {repr(text[:50])}")
        
        return function_titles
    
    def _extract_functions(self) -> List[FunctionInfo]:
        """提取功能列表（包含输入输出要素）
        
        优先级：
        1. 优先从功能说明部分提取功能标题（如 5.2.1、5.2.2 等）
        2. 如果功能清单表存在，用它作为辅助验证
        3. 如果两者不一致，优先使用功能说明中的功能标题
        """
        functions = []
        
        # 1. 查找"5 功能*（A阶段）"章节，忽略之前的内容（非建模需求）
        function_chapter_start = -1
        for i, para in enumerate(self.paragraphs):
            text = para.text.strip()
            if "功能" in text and "（A阶段）" in text and re.match(r'^5\s+功能', text):
                function_chapter_start = i
                parser_logger.debug(f"找到功能章节，索引: {i}, 文本: {repr(text[:100])}")
                break
        
        # 2. 查找"功能说明"部分起始位置（在功能章节之后）
        function_section_start = -1
        search_start = function_chapter_start if function_chapter_start >= 0 else 0
        for i in range(search_start, len(self.paragraphs)):
            text = self.paragraphs[i].text.strip()
            if "功能说明" in text and ("5.2" in text or "（A阶段）" in text):
                function_section_start = i
                break
        
        # 2. 优先从功能说明部分提取功能标题
        function_titles_from_section = []
        if function_section_start >= 0:
            function_titles_from_section = self._extract_function_titles_from_section(function_section_start)
        
        # 3. 从功能清单表提取功能名称（作为辅助）
        function_names_from_table = self._extract_function_list()
        
        # 4. 确定最终使用的功能名称列表
        # 如果功能说明部分有功能标题，优先使用；否则使用功能清单表
        parser_logger.info(f"功能说明部分提取到 {len(function_titles_from_section)} 个功能标题")
        parser_logger.info(f"功能清单表提取到 {len(function_names_from_table)} 个功能名称")
        
        if function_titles_from_section:
            # 使用功能说明部分的功能标题
            function_names = [title for title, _ in function_titles_from_section]
            function_name_to_index = {title: idx for title, idx in function_titles_from_section}
            parser_logger.info(f"✓ 优先使用功能说明部分的功能标题: {function_names}")
        elif function_names_from_table:
            # 回退到功能清单表
            function_names = function_names_from_table
            function_name_to_index = {}
            parser_logger.info(f"✓ 使用功能清单表的功能名称: {function_names}")
        else:
            # 都没有找到
            parser_logger.warning("未找到功能列表：功能说明部分和功能清单表都为空")
            return functions
        
        # 5. 如果功能清单表存在但功能说明部分也有功能标题，验证一致性
        if function_titles_from_section and function_names_from_table:
            if len(function_titles_from_section) != len(function_names_from_table):
                parser_logger.warning(
                    f"功能说明部分的功能数量({len(function_titles_from_section)})与功能清单表的功能数量"
                    f"({len(function_names_from_table)})不一致，优先使用功能说明部分的功能标题"
                )
        
        # 6. 如果使用功能清单表的功能名称，需要建立索引映射
        if not function_name_to_index:
            search_start = max(function_section_start if function_section_start >= 0 else 0, 100)
            search_end = len(self.paragraphs)
            
            for function_name in function_names:
                cleaned_function = re.sub(r"[^\w\u4e00-\u9fa5]", "", function_name)
                found = False
                
                # 优先查找：在功能说明部分内精确匹配功能名称的段落
                if function_section_start >= 0:
                    for i in range(max(function_section_start, 100), search_end):
                        para = self.paragraphs[i]
                        text = para.text.strip()
                        
                        # 精确匹配：段落文本就是功能名称（可能带编号）
                        if function_name == text or (function_name in text and len(text) <= len(function_name) + 10):
                            if "目录" not in text and not re.match(r'^\d+\.\d+', text):
                                function_name_to_index[function_name] = i
                                found = True
                                break
                
                # 如果没找到精确匹配，使用模糊匹配
                if not found:
                    if function_section_start >= 0:
                        for i in range(max(function_section_start, 100), search_end):
                            para = self.paragraphs[i]
                            text = para.text.strip()
                            cleaned_text = re.sub(r"[^\w\u4e00-\u9fa5]", "", text)
                            
                            if (cleaned_function in cleaned_text or cleaned_text in cleaned_function) and len(cleaned_text) >= len(cleaned_function) * 0.7:
                                if ("功能" in text or function_name in text) and "目录" not in text:
                                    if not re.match(r'^\d+\.\d+', text) or len(text) > 50:
                                        function_name_to_index[function_name] = i
                                        found = True
                                        break
                    
                    if not found:
                        for i in range(search_start, search_end):
                            para = self.paragraphs[i]
                            text = para.text.strip()
                            cleaned_text = re.sub(r"[^\w\u4e00-\u9fa5]", "", text)
                            
                            if (cleaned_function in cleaned_text or cleaned_text in cleaned_function) and len(cleaned_text) >= len(cleaned_function) * 0.7:
                                if ("功能" in text or function_name in text) and "目录" not in text:
                                    if not re.match(r'^\d+\.\d+', text) or len(text) > 50:
                                        function_name_to_index[function_name] = i
                                        break
        
        # 7. 为每个功能提取详细输入输出要素（使用缓存的索引，失败时回退到原方法）
        for function_name in function_names:
            function_index = function_name_to_index.get(function_name, -1)
            
            # 如果优化方法找到了索引，使用优化版本；否则回退到原方法
            if function_index >= 0:
                input_elements, output_elements = self._extract_function_input_output_optimized(
                    function_name, function_index, function_section_start
                )
            else:
                # 回退到原来的方法，确保兼容性
                input_elements, output_elements = self._extract_function_input_output(function_name)
            
            function = FunctionInfo(
                name=function_name,
                input_elements=input_elements,
                output_elements=output_elements
            )
            functions.append(function)
        
        return functions
    
    def _extract_function_input_output_optimized(self, function_name: str, function_section_index: int, function_section_start: int) -> Tuple[List[InputElement], List[OutputElement]]:
        """提取指定功能的输入输出要素（优化版本，使用预计算的索引）"""
        input_elements = []
        output_elements = []
        
        if function_section_index < 0:
            return input_elements, output_elements
        
        # 对于非建模需求，只在5.2功能说明部分内查找
        # 查找功能说明部分的结束位置（下一个5.2.x功能或下一个二级标题如"6."）
        end_index = len(self.paragraphs)
        for i in range(function_section_index + 1, len(self.paragraphs)):
            para = self.paragraphs[i]
            text = para.text.strip()
            # 查找下一个5.2.x功能（如"5.2.2"）
            if re.match(r'^5\.2\.\d+', text) and i > function_section_index + 5:
                # 可能是下一个功能，检查是否包含"功能"关键词
                if "功能" in text:
                    end_index = i
                    parser_logger.debug(f"功能 '{function_name}' 搜索范围结束（下一个功能）: {i}, 文本: {repr(text)}")
                    break
            # 查找下一个二级标题（如"6."、"5.3"），表示功能说明部分结束
            if (self._is_heading(para, level=2) or self._is_heading(para, level=1)):
                if re.match(r'^[56]\.', text) or re.match(r'^[56]\.\d+', text):
                    end_index = i
                    parser_logger.debug(f"功能 '{function_name}' 搜索范围结束（下一个章节）: {i}, 文本: {repr(text)}")
                    break
        
        # 搜索范围：从功能位置开始，到功能说明部分结束（或下一个功能）
        search_start = function_section_index
        search_end = end_index
        parser_logger.debug(f"功能 '{function_name}' 搜索范围: {search_start} 到 {search_end}")
        
        # 在功能章节内查找"输入要素"和"输出要素"标记
        input_markers = ["输入要素", "输入要素：", "输入输出要素", "输入要素表"]
        output_markers = ["输出要素", "输出要素：", "输出要素表"]
        
        found_input_marker = False
        found_output_marker = False
        input_marker_index = -1
        output_marker_index = -1
        
        # 先找到"输入要素"和"输出要素"文本的位置，并检查是否"不涉及"
        input_not_involved = False
        output_not_involved = False
        found_input_output_section = False
        found_input_output_elements_section = False
        
        # 步骤1：先查找"输入输出说明"标记，确保我们在正确的章节内
        for i in range(search_start, search_end):
            para = self.paragraphs[i]
            text = para.text.strip()
            
            # 查找"输入输出说明"（可能是 5.2.2.2 输入输出说明 这样的格式）
            if "输入输出说明" in text:
                found_input_output_section = True
                search_start_io = i
                parser_logger.debug(f"找到输入输出说明章节，索引: {i}, 文本: {repr(text)}")
                break
        
        # 步骤2：如果找到了"输入输出说明"章节，在该章节内查找"输入输出要素"子章节
        if found_input_output_section:
            # 先查找"输入输出要素"子章节（如"三、输入输出要素："）
            for i in range(search_start_io + 1, min(search_start_io + 50, search_end)):
                para = self.paragraphs[i]
                text = para.text.strip()
                
                # 查找"输入输出要素"子章节（可能包含"三、"、"三、"等编号）
                if "输入输出要素" in text and ("：" in text or ":" in text):
                    found_input_output_elements_section = True
                    search_start_io_elements = i
                    parser_logger.debug(f"找到输入输出要素子章节，索引: {i}, 文本: {repr(text)}")
                    break
            
            # 步骤3：在"输入输出要素"子章节内查找"输入要素"和"输出要素"标记
            if found_input_output_elements_section:
                # 在输入输出要素子章节内查找，扩大搜索范围到100个段落
                for i in range(search_start_io_elements, min(search_start_io_elements + 100, search_end)):
                    para = self.paragraphs[i]
                    text = para.text.strip()
                
                    # 查找"输入要素"标记（在输入输出要素子章节内）
                    if not found_input_marker:
                        for marker in input_markers:
                            if marker in text and ("输入要素" in marker or marker == "输入要素："):
                                found_input_marker = True
                                input_marker_index = i
                                parser_logger.debug(f"找到输入要素标记，索引: {i}, 文本: {repr(text)}")
                                # 检查"不涉及"：只有"不涉及"直接跟在标记后面（同一段落内），才是不涉及
                                # 例如："输入要素：不涉及" 或 "输入要素： 不涉及"（同一段落）
                                # 如果"输入要素："后面换行了，然后下一行是"不涉及"，那应该不是不涉及
                                
                                # 检查当前段落是否包含"不涉及"（在同一段落内）
                                if "不涉及" in text:
                                    # 检查"不涉及"是否直接跟在"输入要素："后面
                                    marker_pos = text.find("输入要素")
                                    not_involved_pos = text.find("不涉及")
                                    if marker_pos >= 0 and not_involved_pos > marker_pos:
                                        # 如果中间只有冒号、空格、制表符，认为是直接跟在后面
                                        if re.match(r'^输入要素[：:]\s*不涉及', text[marker_pos:]):
                                            input_not_involved = True
                                            parser_logger.debug(f"输入要素标记在同一段落内直接跟着'不涉及'，索引: {i}, 文本: {repr(text[:100])}")
                                # 如果不在同一段落内，不检查下一段落（因为换行后的"不涉及"更可能是其他章节的）
                                break
                
                    # 查找"输出要素"标记（在输入输出要素子章节内）
                    # 注意：不能把"输入输出要素："误认为是"输出要素"标记
                    if not found_output_marker:
                        for marker in output_markers:
                            # 确保是"输出要素"而不是"输入输出要素"
                            if marker in text and ("输出要素" in marker or marker == "输出要素：") and "输入输出要素" not in text:
                                found_output_marker = True
                                output_marker_index = i
                                parser_logger.debug(f"找到输出要素标记，索引: {i}, 文本: {repr(text)}")
                                # 检查"不涉及"：只有"不涉及"直接跟在标记后面（同一段落内），才是不涉及
                                # 例如："输出要素：不涉及" 或 "输出要素： 不涉及"（同一段落）
                                # 如果"输出要素："后面换行了，然后下一行是"不涉及"，那应该不是不涉及
                                
                                # 检查当前段落是否包含"不涉及"（在同一段落内）
                                if "不涉及" in text:
                                    # 检查"不涉及"是否直接跟在"输出要素："后面
                                    marker_pos = text.find("输出要素")
                                    not_involved_pos = text.find("不涉及")
                                    if marker_pos >= 0 and not_involved_pos > marker_pos:
                                        # 如果中间只有冒号、空格、制表符，认为是直接跟在后面
                                        if re.match(r'^输出要素[：:]\s*不涉及', text[marker_pos:]):
                                            output_not_involved = True
                                            parser_logger.debug(f"输出要素标记在同一段落内直接跟着'不涉及'，索引: {i}, 文本: {repr(text[:100])}")
                                # 如果不在同一段落内，不检查下一段落（因为换行后的"不涉及"更可能是其他章节的）
                                break
                    
                    if found_input_marker and found_output_marker:
                        break
        else:
            # 如果没有找到"输入输出说明"，使用原来的逻辑
            for i in range(search_start, search_end):
                para = self.paragraphs[i]
                text = para.text.strip()
                
                # 查找"输入要素"标记
                if not found_input_marker:
                    for marker in input_markers:
                        if marker in text:
                            found_input_marker = True
                            input_marker_index = i
                            # 检查后续段落是否包含"不涉及"
                            for j in range(i + 1, min(i + 4, search_end)):
                                next_text = self.paragraphs[j].text.strip()
                                if any(m in next_text for m in output_markers):
                                    break
                                if "不涉及" in next_text:
                                    input_not_involved = True
                                    break
                            break
                
                # 查找"输出要素"标记
                if not found_output_marker:
                    for marker in output_markers:
                        if marker in text:
                            found_output_marker = True
                            output_marker_index = i
                            # 检查后续段落是否包含"不涉及"
                            # 但需要确保"不涉及"紧跟在输出要素标记后，而不是在其他章节中
                            for j in range(i + 1, min(i + 4, search_end)):
                                next_text = self.paragraphs[j].text.strip()
                                if not next_text:
                                    continue
                                # 如果遇到章节标题（如"三、操作步骤说明："），停止检查
                                if re.match(r'^[一二三四五六七八九十]+[、.]', next_text) and len(next_text) > 5:
                                    # 这是章节标题，停止检查
                                    break
                                # 如果遇到其他输入输出相关的标记，也停止检查
                                if any(m in next_text for m in input_markers + output_markers):
                                    break
                                # 检查"不涉及"：如果后续有章节标题，则"不涉及"不属于输出要素
                                if "不涉及" in next_text and j <= i + 3:
                                    # 检查后续是否有章节标题（最多检查2行）
                                    is_in_section = False
                                    for k in range(j + 1, min(j + 3, search_end)):
                                        follow_text = self.paragraphs[k].text.strip()
                                        if re.match(r'^[一二三四五六七八九十]+[、.]', follow_text) and len(follow_text) > 5:
                                            # "不涉及"后面有章节标题，说明"不涉及"属于该章节，不属于输出要素
                                            is_in_section = True
                                            break
                                    if not is_in_section:
                                        # "不涉及"后面没有章节标题，说明它属于输出要素
                                        output_not_involved = True
                                        break
                            break
                
                if found_input_marker and found_output_marker:
                    break
        
        # 智能表格定位：在功能章节范围内查找表格
        if found_input_marker and not input_not_involved:
            # 优先在功能章节范围内查找标记后最近的输入要素表
            if input_marker_index >= 0:
                # 使用功能章节范围（从输入要素标记到章节结束）查找表格
                input_elements = self._find_table_after_marker_in_range(
                    input_marker_index, end_index, is_input=True, is_document_end=False, function_name=function_name
                )
            # 如果没找到，再查找所有未使用的表格（作为回退策略）
            if not input_elements:
                input_elements = self._search_all_unused_tables(is_input=True)
        elif input_not_involved:
            input_elements = []
        
        if found_output_marker and not output_not_involved:
            # 优先在"输入输出说明"章节范围内查找标记后最近的输出要素表
            if output_marker_index >= 0:
                # 查找"输入输出说明"章节的结束位置（通常是"操作步骤说明"或下一个子章节）
                io_section_end = end_index  # 默认使用功能章节结束位置
                if found_input_output_section:
                    # 在"输入输出说明"章节后查找下一个子章节标题
                    for i in range(search_start_io + 1, min(search_start_io + 50, search_end)):
                        para = self.paragraphs[i]
                        text = para.text.strip()
                        # 查找下一个子章节标题（如"操作步骤说明"、"三、"等）
                        if (re.match(r'^[一二三四五六七八九十]+[、.]', text) and len(text) > 5) or \
                           ("操作步骤说明" in text) or \
                           (self._is_heading(para, level=3) and i > search_start_io + 10):
                            io_section_end = i
                            break
                
                # 使用"输入输出说明"章节范围（从输出要素标记到该章节结束）查找表格
                output_elements = self._find_table_after_marker_in_range(
                    output_marker_index, io_section_end, is_input=False, is_document_end=False, function_name=function_name
                )
            # 如果没找到，再查找所有未使用的表格（作为回退策略）
            if not output_elements:
                output_elements = self._search_all_unused_tables(is_input=False)
        elif output_not_involved:
            output_elements = []
        
        return input_elements, output_elements
    
    def _extract_function_input_output(self, function_name: str) -> Tuple[List[InputElement], List[OutputElement]]:
        """提取指定功能的输入输出要素"""
        input_elements = []
        output_elements = []
        
        # 查找功能名称所在位置（可能在标题中，也可能在普通段落中）
        # 优先查找功能说明部分（5.2）下的功能名称
        function_section_index = -1
        
        # 清理功能名称用于匹配（去除标点符号）
        cleaned_function = re.sub(r"[^\w\u4e00-\u9fa5]", "", function_name)
        
        # 先查找"功能说明"部分
        function_section_start = -1
        for i, para in enumerate(self.paragraphs):
            text = para.text.strip()
            if "功能说明" in text and ("5.2" in text or "（A阶段）" in text):
                function_section_start = i
                break
        
        # 在功能说明部分查找功能名称（优先匹配精确的功能名称段落）
        search_start = function_section_start if function_section_start >= 0 else 0
        search_end = len(self.paragraphs)
        
        # 优先查找：在功能说明部分内精确匹配功能名称的段落
        if function_section_start >= 0:
            # 在功能说明部分内查找（跳过目录部分，通常目录在前100个段落）
            for i in range(max(function_section_start, 100), search_end):
                para = self.paragraphs[i]
                text = para.text.strip()
                
                # 精确匹配：段落文本就是功能名称（可能带编号）
                if function_name == text or (function_name in text and len(text) <= len(function_name) + 10):
                    # 排除目录和编号行
                    if "目录" not in text and not re.match(r'^\d+\.\d+', text):
                        function_section_index = i
                        break
        
        # 如果没找到精确匹配，使用模糊匹配（优先在功能说明部分内）
        if function_section_index < 0:
            # 先在功能说明部分内查找
            if function_section_start >= 0:
                for i in range(max(function_section_start, 100), search_end):
                    para = self.paragraphs[i]
                    text = para.text.strip()
                    
                    # 检查是否匹配功能名称（可能是标题或普通段落）
                    cleaned_text = re.sub(r"[^\w\u4e00-\u9fa5]", "", text)
                    
                    # 匹配逻辑：功能名称完全匹配，或者功能名称包含在文本中
                    if (cleaned_function in cleaned_text or cleaned_text in cleaned_function) and len(cleaned_text) >= len(cleaned_function) * 0.7:
                        # 确保不是在目录或其他不相关的地方
                        if ("功能" in text or function_name in text) and "目录" not in text:
                            # 排除目录行（通常包含页码）
                            if not re.match(r'^\d+\.\d+', text) or len(text) > 50:
                                function_section_index = i
                                break
            
            # 如果还没找到，在整个文档中查找
            if function_section_index < 0:
                for i in range(search_start, search_end):
                    para = self.paragraphs[i]
                    text = para.text.strip()
                    
                    cleaned_text = re.sub(r"[^\w\u4e00-\u9fa5]", "", text)
                    
                    if (cleaned_function in cleaned_text or cleaned_text in cleaned_function) and len(cleaned_text) >= len(cleaned_function) * 0.7:
                        if ("功能" in text or function_name in text) and "目录" not in text:
                            # 排除目录行
                            if not re.match(r'^\d+\.\d+', text) or len(text) > 50:
                                function_section_index = i
                                break
        
        if function_section_index < 0:
            return input_elements, output_elements
        
        # 在功能章节内查找输入输出要素
        # 查找结束位置（下一个三级标题、二级标题或一级标题）
        end_index = len(self.paragraphs)
        for i in range(function_section_index + 1, len(self.paragraphs)):
            para = self.paragraphs[i]
            if (self._is_heading(para, level=3) or 
                self._is_heading(para, level=2) or 
                self._is_heading(para, level=1)):
                end_index = i
                break
        
        # 扩大搜索范围：从功能章节开始，向后搜索
        search_start = max(0, function_section_index - 10)
        search_end = min(len(self.paragraphs), function_section_index + 200)
        
        # 在功能章节内查找"输入要素"和"输出要素"标记
        input_markers = ["输入要素", "输入要素：", "输入输出要素", "输入要素表"]
        output_markers = ["输出要素", "输出要素：", "输出要素表"]
        
        found_input_marker = False
        found_output_marker = False
        input_marker_index = -1
        output_marker_index = -1
        
        # 先找到"输入要素"和"输出要素"文本的位置，并检查是否"不涉及"
        # 需要找到该功能章节内的第一个"输入输出要素"标记（在"输入输出说明"下）
        input_not_involved = False
        output_not_involved = False
        found_input_output_section = False
        
        # 先查找"输入输出说明"标记，确保我们在正确的章节内
        for i in range(search_start, search_end):
            para = self.paragraphs[i]
            text = para.text.strip()
            
            # 查找"输入输出说明"或"输入输出要素"
            if "输入输出说明" in text or ("输入输出要素" in text and "：" in text):
                found_input_output_section = True
                # 从"输入输出说明"开始查找输入输出要素标记
                search_start_io = i
                break
        
        # 如果找到了"输入输出说明"章节，在该章节内查找
        if found_input_output_section:
            for i in range(search_start_io, min(search_start_io + 20, search_end)):
                para = self.paragraphs[i]
                text = para.text.strip()
                
                # 查找"输入要素"标记（必须在"输入输出说明"章节内）
                if not found_input_marker:
                    for marker in input_markers:
                        if marker in text and ("输入要素" in marker or marker == "输入要素："):
                            found_input_marker = True
                            input_marker_index = i
                            # 检查后续段落（最多5行）是否包含"不涉及"
                            # 需要跳过空行，直到找到下一个标记或"不涉及"
                            for j in range(i + 1, min(i + 6, search_end)):
                                next_text = self.paragraphs[j].text.strip()
                                # 跳过空行
                                if not next_text:
                                    continue
                                # 如果遇到下一个标记（如"输出要素"），停止检查
                                if any(m in next_text for m in output_markers):
                                    break
                                # 如果遇到下一个章节标记（如"三、"），停止检查
                                if re.match(r'^[一二三四五六七八九十]+、', next_text):
                                    break
                                # 检查是否包含"不涉及"
                                if "不涉及" in next_text:
                                    input_not_involved = True
                                    break
                            break
                
                # 查找"输出要素"标记（必须在"输入输出说明"章节内）
                if not found_output_marker:
                    for marker in output_markers:
                        if marker in text and ("输出要素" in marker or marker == "输出要素："):
                            found_output_marker = True
                            output_marker_index = i
                            # 检查后续段落（最多5行）是否包含"不涉及"
                            # 但需要确保"不涉及"紧跟在输出要素标记后，而不是在其他章节中
                            for j in range(i + 1, min(i + 6, search_end)):
                                next_text = self.paragraphs[j].text.strip()
                                # 跳过空行
                                if not next_text:
                                    continue
                                # 如果遇到章节标题（如"三、操作步骤说明："），停止检查
                                if re.match(r'^[一二三四五六七八九十]+[、.]', next_text) and len(next_text) > 5:
                                    # 这是章节标题，停止检查
                                    break
                                # 如果遇到其他输入输出相关的标记，也停止检查
                                if any(m in next_text for m in input_markers + output_markers):
                                    break
                                # 检查"不涉及"：如果后续有章节标题，则"不涉及"不属于输出要素
                                if "不涉及" in next_text and j <= i + 3:
                                    # 检查后续是否有章节标题（最多检查2行）
                                    is_in_section = False
                                    for k in range(j + 1, min(j + 3, search_end)):
                                        follow_text = self.paragraphs[k].text.strip()
                                        if re.match(r'^[一二三四五六七八九十]+[、.]', follow_text) and len(follow_text) > 5:
                                            # "不涉及"后面有章节标题，说明"不涉及"属于该章节，不属于输出要素
                                            is_in_section = True
                                            break
                                    if not is_in_section:
                                        # "不涉及"后面没有章节标题，说明它属于输出要素
                                        output_not_involved = True
                                        break
                            break
                
                if found_input_marker and found_output_marker:
                    break
        else:
            # 如果没有找到"输入输出说明"，使用原来的逻辑
            for i in range(search_start, search_end):
                para = self.paragraphs[i]
                text = para.text.strip()
                
                # 查找"输入要素"标记
                if not found_input_marker:
                    for marker in input_markers:
                        if marker in text:
                            found_input_marker = True
                            input_marker_index = i
                            # 检查后续段落是否包含"不涉及"
                            for j in range(i + 1, min(i + 4, search_end)):
                                next_text = self.paragraphs[j].text.strip()
                                if any(m in next_text for m in output_markers):
                                    break
                                if "不涉及" in next_text:
                                    input_not_involved = True
                                    break
                            break
                
                # 查找"输出要素"标记
                if not found_output_marker:
                    for marker in output_markers:
                        if marker in text:
                            found_output_marker = True
                            output_marker_index = i
                            # 检查后续段落是否包含"不涉及"
                            # 但需要确保"不涉及"紧跟在输出要素标记后，而不是在其他章节中
                            for j in range(i + 1, min(i + 4, search_end)):
                                next_text = self.paragraphs[j].text.strip()
                                if not next_text:
                                    continue
                                # 如果遇到章节标题（如"三、操作步骤说明："），停止检查
                                if re.match(r'^[一二三四五六七八九十]+[、.]', next_text) and len(next_text) > 5:
                                    # 这是章节标题，停止检查
                                    break
                                # 如果遇到其他输入输出相关的标记，也停止检查
                                if any(m in next_text for m in input_markers + output_markers):
                                    break
                                # 检查"不涉及"：如果后续有章节标题，则"不涉及"不属于输出要素
                                if "不涉及" in next_text and j <= i + 3:
                                    # 检查后续是否有章节标题（最多检查2行）
                                    is_in_section = False
                                    for k in range(j + 1, min(j + 3, search_end)):
                                        follow_text = self.paragraphs[k].text.strip()
                                        if re.match(r'^[一二三四五六七八九十]+[、.]', follow_text) and len(follow_text) > 5:
                                            # "不涉及"后面有章节标题，说明"不涉及"属于该章节，不属于输出要素
                                            is_in_section = True
                                            break
                                    if not is_in_section:
                                        # "不涉及"后面没有章节标题，说明它属于输出要素
                                        output_not_involved = True
                                        break
                            break
                
                if found_input_marker and found_output_marker:
                    break
        
        # 智能表格定位：按顺序查找未使用的表格
        # 查找输入要素表（只有在不是"不涉及"的情况下才查找）
        if found_input_marker and not input_not_involved:
            # 方法1：直接按顺序查找第一个未使用的输入要素表
            input_elements = self._search_all_unused_tables(is_input=True)
            
            # 方法2：如果没找到，尝试查找标记后最近的一个输入要素表（即使已使用）
            # 这对于"优惠利息查询"等后面功能很重要，因为它们的表格可能在已使用的表格之后
            if not input_elements and input_marker_index >= 0:
                # 查找标记后最近的输入要素表
                input_elements = self._find_nearest_table_after_marker(
                    input_marker_index, is_input=True
                )
        elif input_not_involved:
            # 如果标记为"不涉及"，不提取输入要素
            input_elements = []
        
        # 查找输出要素表（只有在不是"不涉及"的情况下才查找）
        if found_output_marker and not output_not_involved:
            # 方法1：直接按顺序查找第一个未使用的输出要素表
            output_elements = self._search_all_unused_tables(is_input=False)
            
            # 方法2：如果没找到，尝试查找标记后最近的一个输出要素表（即使已使用）
            if not output_elements and output_marker_index >= 0:
                # 查找标记后最近的输出要素表
                output_elements = self._find_nearest_table_after_marker(
                    output_marker_index, is_input=False
                )
        elif output_not_involved:
            # 如果标记为"不涉及"，不提取输出要素
            output_elements = []
        
        return input_elements, output_elements

    def _strip_phase_suffix(self, text: str) -> str:
        """移除标题中可能存在的阶段后缀，如'（A阶段）'、'（A阶段、B阶段）'"""
        return re.sub(r"\*?（[^）]*阶段[^）]*）\s*$", "", text).strip()
