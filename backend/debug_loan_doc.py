"""
调试发放贷款文档结构
"""
import sys
from pathlib import Path
from docx import Document

sys.path.insert(0, str(Path(__file__).parent))

def main():
    doc_path = Path(__file__).parent.parent / "tmp" / "新一代信贷系统建设项目_发放贷款用例_贷款核算组_V0.9_20251113(1).docx"
    
    doc = Document(str(doc_path))
    
    print("=" * 80)
    print("查找任务规则说明和步骤信息：")
    print("=" * 80)
    
    found_rule_section = False
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name
        
        # 查找任务规则说明
        if "任务规则说明" in text and "（A阶段、B阶段）" in text:
            found_rule_section = True
            print(f"\n[{i:3d}] [{style:15s}] {text}")
            continue
        
        # 在任务规则说明部分查找步骤
        if found_rule_section:
            # 查找组件（二级标题）
            if style == 'Heading 2' and "（A阶段、B阶段）" in text:
                print(f"\n[{i:3d}] [{style:15s}] 组件: {text}")
            
            # 查找任务（三级标题）
            elif style == 'Heading 3' and "（A阶段、B阶段）" in text:
                print(f"[{i:3d}] [{style:15s}] 任务: {text}")
            
            # 查找步骤（四级标题）
            elif style == 'Heading 4' and "（A阶段、B阶段）" in text:
                print(f"[{i:3d}] [{style:15s}] 步骤: {text}")
                
                # 查找该步骤下的输入输出
                for j in range(i + 1, min(i + 50, len(doc.paragraphs))):
                    next_para = doc.paragraphs[j]
                    next_text = next_para.text.strip()
                    next_style = next_para.style.name
                    
                    # 如果遇到下一个四级标题，停止
                    if next_style == 'Heading 4' and "（A阶段、B阶段）" in next_text:
                        break
                    
                    # 如果遇到三级标题，停止
                    if next_style == 'Heading 3':
                        break
                    
                    # 如果遇到二级标题，停止
                    if next_style == 'Heading 2':
                        break
                    
                    # 查找输入输出标题
                    if "输入输出" in next_text and "（A阶段、B阶段）" in next_text:
                        print(f"    [{j:3d}] [{next_style:15s}] 找到输入输出标题")
                        # 查找表格
                        for table_idx, table in enumerate(doc.tables):
                            if len(table.rows) > 0:
                                first_row = ' '.join([cell.text.strip() for cell in table.rows[0].cells])
                                if "输入" in first_row and "字段名称" in first_row:
                                    print(f"        找到输入要素表（表格{table_idx+1}），行数: {len(table.rows)}")
                                elif "输出" in first_row and "字段名称" in first_row:
                                    print(f"        找到输出要素表（表格{table_idx+1}），行数: {len(table.rows)}")
                        break
            
            # 如果遇到下一个一级标题，停止搜索
            if style == 'Heading 1' and "任务规则说明" not in text:
                break

if __name__ == "__main__":
    main()

