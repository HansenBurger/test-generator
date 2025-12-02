"""
调试文档结构
"""
import sys
from pathlib import Path
from docx import Document

sys.path.insert(0, str(Path(__file__).parent))

def main():
    doc_path = Path(__file__).parent.parent / "tmp" / "新一代信贷系统建设项目_管理特色互联网贷款账单用例_贷款核算组_V0.9_20251028.docx"
    
    doc = Document(str(doc_path))
    
    print("=" * 80)
    print("文档段落信息（前100个）：")
    print("=" * 80)
    
    for i, para in enumerate(doc.paragraphs[:100]):
        text = para.text.strip()
        if text:
            style = para.style.name
            print(f"[{i:3d}] [{style:15s}] {text[:100]}")
    
    print("\n" + "=" * 80)
    print("文档表格信息：")
    print("=" * 80)
    
    for i, table in enumerate(doc.tables):
        print(f"\n表格 {i+1}:")
        if len(table.rows) > 0:
            # 打印表头
            header = [cell.text.strip() for cell in table.rows[0].cells]
            print(f"  表头: {header}")
            
            # 打印前3行数据
            for j, row in enumerate(table.rows[1:4]):
                if j >= 3:
                    break
                cells = [cell.text.strip() for cell in row.cells]
                print(f"  行{j+1}: {cells}")

if __name__ == "__main__":
    main()

